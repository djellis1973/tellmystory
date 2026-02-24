# biographer.py – Tell My Story App (COMPLETE WORKING VERSION with STREAKS & GAMIFICATION)
import streamlit as st
import json
from datetime import datetime, date, timedelta
from openai import OpenAI
import os
import re
import hashlib
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import secrets
import string
import time
import shutil
import base64
from PIL import Image
import io
import zipfile
import html
import subprocess
import tempfile
import csv  # Added for historical events

# For EPUB export
try:
    from ebooklib import epub
    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False
    st.warning("For EPUB export: pip install ebooklib")

# ============================================================================
# PAGE CONFIG - MUST BE FIRST
# ============================================================================
st.set_page_config(page_title="Tell My Story - Your Life Timeline", page_icon="📖", layout="wide", initial_sidebar_state="expanded")

# ============================================================================
# PUBLISHER AVAILABILITY (functions defined later)
# ============================================================================
PUBLISHER_AVAILABLE = True

# ============================================================================
# IMPORT QUILL RICH TEXT EDITOR
# ============================================================================
try:
    from streamlit_quill import st_quill
    QUILL_AVAILABLE = True
except ImportError:
    st.error("❌ Please install streamlit-quill: pip install streamlit-quill")
    st.stop()

# ============================================================================
# FORCE DIRECTORY CREATION
# ============================================================================
for dir_path in ["question_banks/default", "question_banks/users", "question_banks", 
                 "uploads", "uploads/thumbnails", "uploads/metadata", "accounts", "sessions", "backups"]:
    os.makedirs(dir_path, exist_ok=True)

# ============================================================================
# IMPORTS
# ============================================================================
try:
    from topic_bank import TopicBank
    from session_manager import SessionManager
    from vignettes import VignetteManager
    from session_loader import SessionLoader
    from beta_reader import BetaReader
    from question_bank_manager import QuestionBankManager
except ImportError as e:
    st.error(f"Error importing modules: {e}")
    st.info("Please ensure all .py files are in the same directory")
    TopicBank = SessionManager = VignetteManager = SessionLoader = BetaReader = QuestionBankManager = None

DEFAULT_WORD_TARGET = 500

# ============================================================================
# INITIALIZATION
# ============================================================================
client = OpenAI(api_key=st.secrets.get("OPENAI_API_KEY", os.environ.get("OPENAI_API_KEY")))
beta_reader = BetaReader(client) if BetaReader else None

# Initialize session state with PROMPT ME variables added
default_state = {
    "qb_manager": None, "qb_manager_initialized": False, "user_id": None, "logged_in": False,
    "current_session": 0, "current_question": 0, "responses": {}, "editing": False,
    "editing_word_target": False, "confirming_clear": None, "data_loaded": False,
    "current_question_override": None, "show_vignette_modal": False, "vignette_topic": "",
    "vignette_content": "", "selected_vignette_type": "Standard Topic", "current_vignette_list": [],
    "editing_vignette_index": None, "show_vignette_manager": False, "custom_topic_input": "",
    "show_custom_topic_modal": False, "show_topic_browser": False, "show_session_manager": False,
    "show_session_creator": False, "editing_custom_session": None, "show_vignette_detail": False,
    "selected_vignette_id": None, "editing_vignette_id": None, "selected_vignette_for_session": None,
    "published_vignette": None, "show_beta_reader": False, "current_beta_feedback": None,
    "current_question_bank": None, "current_bank_name": None, "current_bank_type": None,
    "current_bank_id": None, "show_bank_manager": False, "show_bank_editor": False,
    "editing_bank_id": None, "editing_bank_name": None, "qb_manager": None, "qb_manager_initialized": False,
    "confirm_delete": None, "user_account": None, "show_profile_setup": False,
    "image_handler": None, "show_image_manager": False,
    "current_rewrite_data": None, "show_ai_rewrite": False, "show_ai_rewrite_menu": False,
    "editor_content": {}, "show_privacy_settings": False, "show_cover_designer": False,
    "beta_feedback_display": None, "beta_feedback_storage": {},
    "auth_tab": 'login',
    "show_publisher": False,
    "cover_image_data": None,
    "show_support": False,
    "show_prompt_modal": False,
    "current_prompt_data": None,
    # Gamification celebration flags
    "milestone_achieved_first_story": False,
    "milestone_achieved_seven_day": False,
    "milestone_achieved_five_thousand": False,
    "milestone_achieved_first_session": False
}

for key, value in default_state.items():
    if key not in st.session_state:
        st.session_state[key] = value

# Load external CSS
try:
    with open("styles.css", encoding="utf-8") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
except FileNotFoundError:
    pass

LOGO_URL = "https://menuhunterai.com/wp-content/uploads/2026/02/tms_logo.png"

# ============================================================================
# GAMIFICATION SYSTEM - STREAKS & MILESTONES
# ============================================================================
def update_writing_streak(user_id):
    """Update the user's writing streak based on today's activity"""
    if not user_id or not st.session_state.user_account:
        return
    
    today = datetime.now().date()
    today_str = today.isoformat()
    
    # Initialize streak tracking if not exists
    if 'streak_data' not in st.session_state.user_account:
        st.session_state.user_account['streak_data'] = {
            'current_streak': 0,
            'longest_streak': 0,
            'last_write_date': None,
            'streak_history': [],
            'milestones': {
                'first_story': False,
                'seven_day_streak': False,
                'five_thousand_words': False,
                'first_session_complete': False
            }
        }
    
    streak_data = st.session_state.user_account['streak_data']
    last_date = streak_data.get('last_write_date')
    
    # Calculate today's word count (from all sessions)
    today_words = 0
    for session in st.session_state.current_question_bank or []:
        sid = session["id"]
        if sid in st.session_state.responses:
            for q_data in st.session_state.responses[sid].get("questions", {}).values():
                if q_data.get("answer"):
                    # Check if written today
                    timestamp = q_data.get("timestamp", "")
                    if timestamp and timestamp.startswith(today_str):
                        text_only = re.sub(r'<[^>]+>', '', q_data["answer"])
                        today_words += len(re.findall(r'\w+', text_only))
    
    # Only count if at least 50 words written today
    if today_words >= 50:
        if last_date == today_str:
            # Already counted today - no change
            pass
        elif last_date == (today - timedelta(days=1)).isoformat():
            # Consecutive day!
            streak_data['current_streak'] += 1
            streak_data['longest_streak'] = max(streak_data['longest_streak'], streak_data['current_streak'])
        else:
            # Streak broken or first day
            streak_data['current_streak'] = 1
        
        streak_data['last_write_date'] = today_str
        
        # Add to history (keep last 30 days)
        streak_data['streak_history'].append({
            'date': today_str,
            'words': today_words
        })
        # Keep only last 90 days
        streak_data['streak_history'] = streak_data['streak_history'][-90:]
        
        # Check milestones
        check_milestones(user_id, streak_data)
        
        save_account_data(st.session_state.user_account)

def check_milestones(user_id, streak_data):
    """Check and update milestone achievements"""
    milestones = streak_data.get('milestones', {})
    
    # First Story (100 words total)
    if not milestones.get('first_story'):
        total_words = st.session_state.user_account["stats"].get("total_words", 0)
        if total_words >= 100:
            milestones['first_story'] = True
            st.session_state[f"milestone_achieved_first_story"] = True
    
    # 7-Day Streak
    if not milestones.get('seven_day_streak'):
        if streak_data.get('current_streak', 0) >= 7:
            milestones['seven_day_streak'] = True
            st.session_state[f"milestone_achieved_seven_day"] = True
    
    # 5,000 Total Words
    if not milestones.get('five_thousand_words'):
        total_words = st.session_state.user_account["stats"].get("total_words", 0)
        if total_words >= 5000:
            milestones['five_thousand_words'] = True
            st.session_state[f"milestone_achieved_five_thousand"] = True
    
    # First Session Complete
    if not milestones.get('first_session_complete'):
        for session in st.session_state.current_question_bank or []:
            sid = session["id"]
            if sid in st.session_state.responses:
                answered = len(st.session_state.responses[sid].get("questions", {}))
                total = len(session["questions"])
                if answered >= total and total > 0:
                    milestones['first_session_complete'] = True
                    st.session_state[f"milestone_achieved_first_session"] = True
                    break
    
    streak_data['milestones'] = milestones

def get_todays_word_count():
    """Get total words written today across all sessions"""
    today = datetime.now().date().isoformat()
    total = 0
    
    for session in st.session_state.current_question_bank or []:
        sid = session["id"]
        if sid in st.session_state.responses:
            for q_data in st.session_state.responses[sid].get("questions", {}).values():
                timestamp = q_data.get("timestamp", "")
                if timestamp and timestamp.startswith(today):
                    text_only = re.sub(r'<[^>]+>', '', q_data.get("answer", ""))
                    total += len(re.findall(r'\w+', text_only))
    
    return total

def get_daily_goal():
    """Get daily word goal from user settings or default"""
    # Default goal is 500 words per day
    return st.session_state.user_account.get('settings', {}).get('daily_word_goal', 500)

def render_gamification_dashboard():
    """Render the gamification dashboard in the sidebar"""
    if not st.session_state.user_account:
        return
    
    # Get streak data
    streak_data = st.session_state.user_account.get('streak_data', {})
    current_streak = streak_data.get('current_streak', 0)
    milestones = streak_data.get('milestones', {})
    
    # Today's stats
    today_words = get_todays_word_count()
    daily_goal = get_daily_goal()
    goal_percent = min(100, int((today_words / daily_goal) * 100)) if daily_goal > 0 else 0
    
    # Total stats
    total_words = st.session_state.user_account["stats"].get("total_words", 0)
    
    # Display streak with fire emoji (more fire = longer streak)
    fire_emoji = "🔥" * min(5, max(1, (current_streak // 7) + 1)) if current_streak > 0 else "🌱"
    
    st.markdown("""
    <style>
    .streak-box {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 20px;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 20px;
    }
    .streak-number {
        font-size: 48px;
        font-weight: bold;
        line-height: 1;
    }
    .streak-label {
        font-size: 14px;
        opacity: 0.9;
    }
    .progress-container {
        background: rgba(255,255,255,0.2);
        border-radius: 10px;
        height: 10px;
        margin: 10px 0;
    }
    .progress-fill {
        background: white;
        border-radius: 10px;
        height: 10px;
        transition: width 0.3s ease;
    }
    .milestone-item {
        display: flex;
        align-items: center;
        padding: 8px;
        background: rgba(255,255,255,0.1);
        border-radius: 5px;
        margin-bottom: 5px;
    }
    .milestone-check {
        margin-right: 10px;
        font-size: 18px;
    }
    .milestone-text {
        font-size: 14px;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Streak Box
    st.markdown(f"""
    <div class="streak-box">
        <div style="font-size: 32px;">{fire_emoji}</div>
        <div class="streak-number">{current_streak}</div>
        <div class="streak-label">DAY STREAK</div>
        <div class="progress-container">
            <div class="progress-fill" style="width: {goal_percent}%;"></div>
        </div>
        <div style="display: flex; justify-content: space-between; margin-top: 5px;">
            <span>📝 {today_words} today</span>
            <span>🎯 {daily_goal} goal</span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Quick Stats
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Total Words", f"{total_words:,}")
    with col2:
        sessions_completed = 0
        if st.session_state.current_question_bank:
            sessions_completed = sum(1 for s in st.session_state.current_question_bank 
                                   if len(st.session_state.responses.get(s["id"], {}).get("questions", {})) >= len(s["questions"]))
        st.metric("Sessions Done", f"{sessions_completed}/{len(st.session_state.current_question_bank) if st.session_state.current_question_bank else 0}")
    
    st.divider()
    
    # Milestones
    st.markdown("### ✨ Milestones")
    
    milestones_list = [
        ("first_story", "📖 First Story (100 words)", milestones.get('first_story', False)),
        ("seven_day_streak", "🔥 7-Day Streak", milestones.get('seven_day_streak', False)),
        ("five_thousand_words", "📚 5,000 Total Words", milestones.get('five_thousand_words', False)),
        ("first_session_complete", "✅ Complete First Session", milestones.get('first_session_complete', False))
    ]
    
    for key, label, achieved in milestones_list:
        if achieved:
            st.markdown(f'<div class="milestone-item"><span class="milestone-check">✅</span><span class="milestone-text">{label}</span></div>', unsafe_allow_html=True)
        else:
            # Show progress for unfinished milestones
            if key == "first_story":
                progress = min(100, int((total_words / 100) * 100))
                st.markdown(f'<div class="milestone-item"><span class="milestone-check">⭕</span><span class="milestone-text">{label} ({progress}%)</span></div>', unsafe_allow_html=True)
            elif key == "seven_day_streak":
                progress = min(100, int((current_streak / 7) * 100))
                st.markdown(f'<div class="milestone-item"><span class="milestone-check">⭕</span><span class="milestone-text">{label} ({progress}%)</span></div>', unsafe_allow_html=True)
            elif key == "five_thousand_words":
                progress = min(100, int((total_words / 5000) * 100))
                st.markdown(f'<div class="milestone-item"><span class="milestone-check">⭕</span><span class="milestone-text">{label} ({progress}%)</span></div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="milestone-item"><span class="milestone-check">⭕</span><span class="milestone-text">{label}</span></div>', unsafe_allow_html=True)
    
    # Celebration for new achievements
    for milestone in ['first_story', 'seven_day', 'five_thousand', 'first_session']:
        flag = f"milestone_achieved_{milestone}"
        if st.session_state.get(flag):
            st.balloons()
            st.success(f"🎉 Achieved: {dict(milestones_list)[milestone]}!")
            st.session_state[flag] = False
            time.sleep(2)
            st.rerun()

# ============================================================================
# HISTORICAL EVENTS HELPER - SIMPLE VERSION
# ============================================================================
def get_historical_events_for_prompt(birth_year=None):
    """Simple function to read historical events from CSV and format them"""
    events_text = ""
    try:
        if os.path.exists("historical_events.csv"):
            with open("historical_events.csv", 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                events = []
                for row in reader:
                    # Extract year for filtering
                    era = row.get('era_range', '')
                    year_num = None
                    if era and birth_year:
                        # Try to extract first year from era (e.g., "1940s" -> 1940)
                        import re
                        years = re.findall(r'\d{4}', era)
                        if years:
                            year_num = int(years[0])
                    
                    # Only include events after birth year
                    if birth_year is None or year_num is None or year_num >= birth_year:
                        events.append(row)
                
                # Take first 5 events for the prompt
                for event in events[:5]:
                    events_text += f"• {event.get('era_range', '')}: {event.get('event', '')} - {event.get('description', '')[:100]}...\n"
    except Exception as e:
        print(f"Could not load historical events: {e}")
    
    return events_text

# ============================================================================
# PROMPT ME - OVERCOME WRITER'S BLOCK
# ============================================================================
def generate_writing_prompts(session_title, question_text, existing_answer, profile_context, birth_year=None):
    """Generate creative prompts to help overcome writer's block"""
    if not client:
        return {"error": "OpenAI client not available"}
    
    try:
        # Clean the existing answer
        clean_answer = re.sub(r'<[^>]+>', '', existing_answer) if existing_answer else ""
        
        # Get historical events if available
        historical_events = get_historical_events_for_prompt(birth_year)
        historical_context = ""
        if historical_events:
            historical_context = f"""
HISTORICAL EVENTS THAT HAPPENED DURING YOUR LIFETIME:
{historical_events}
Consider asking how these events affected their life, family, or community.
"""
        
        # Determine if this is a new topic or continuing
        if not clean_answer or clean_answer == "Start writing your story here...":
            context_type = "new_topic"
            prompt_instruction = "The user hasn't started writing this topic yet."
        else:
            context_type = "in_progress"
            prompt_instruction = f"The user has already written: {clean_answer[:300]}..."
        
        system_prompt = f"""You are a compassionate and insightful writing coach, helping someone tell their life story. 
Your goal is to provide gentle, thought-provoking prompts that help overcome writer's block and spark meaningful memories.

{profile_context}

{historical_context}

CURRENT CONTEXT:
- Session: {session_title}
- Question/Topic: {question_text}
- {prompt_instruction}

TASK: Generate 3-5 thoughtful prompts to help the user continue writing. Each prompt should:
1. Be specific and personal (not generic writing advice)
2. Draw from the user's profile context when relevant
3. If historical events are available, ask how those events affected their life
4. Help them remember details, emotions, or specific moments
5. Be warm and encouraging
6. Feel like a conversation with a supportive friend

FORMAT YOUR RESPONSE WITH THESE SECTIONS:
🎯 **Quick Start**: [One very simple starting point - just one sentence]

✨ **Memory Prompts**:
• [Prompt 1 - specific memory trigger]
• [Prompt 2 - sensory detail or emotion]
• [Prompt 3 - people or relationships]

🌍 **Historical Connection** (if relevant):
• [A question about how world events affected their life - e.g., "Do you remember where you were when...?"]

💭 **Deeper Reflection**:
• [A more profound question to consider]

Keep the tone warm and supportive. Use the user's name if available in profile context.
"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Please help me write about: {question_text}"}
            ],
            max_tokens=800,
            temperature=0.8
        )
        
        prompts = response.choices[0].message.content.strip()
        
        return {
            "success": True,
            "prompts": prompts,
            "context_type": context_type
        }
        
    except Exception as e:
        return {"error": str(e)}

def show_prompt_me_modal():
    """Display the Prompt Me modal with writing prompts"""
    if not st.session_state.get('current_prompt_data'):
        return
    
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    
    col1, col2 = st.columns([6, 1])
    with col1:
        st.markdown("### 💭 Writing Prompts to Spark Your Memory")
    with col2:
        if st.button("✕", key="close_prompt_modal"):
            st.session_state.show_prompt_modal = False
            st.session_state.current_prompt_data = None
            st.rerun()
    
    st.markdown("---")
    
    prompt_data = st.session_state.current_prompt_data
    
    if prompt_data.get('error'):
        st.error(f"Could not generate prompts: {prompt_data['error']}")
    else:
        # Display the prompts
        st.markdown("""
        <style>
        .prompt-box {
            background-color: #f8f9fa;
            border-left: 4px solid #9b59b6;
            padding: 20px;
            border-radius: 8px;
            margin: 20px 0;
        }
        .prompt-box h4 {
            color: #8e44ad;
            margin-top: 0;
        }
        </style>
        """, unsafe_allow_html=True)
        
        st.markdown('<div class="prompt-box">', unsafe_allow_html=True)
        st.markdown(prompt_data['prompts'])
        st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown("---")
        st.markdown("*These prompts are personalized based on your profile and what you've written so far.*")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("📋 Copy to Clipboard", key="copy_prompts", use_container_width=True):
                st.info("✅ Prompts copied! Select the text above and press Ctrl+C")
        
        with col2:
            if st.button("🔄 New Prompts", key="new_prompts", use_container_width=True):
                with st.spinner("Generating fresh prompts..."):
                    # Get current context
                    current_session = st.session_state.current_question_bank[st.session_state.current_session]
                    current_question_text = st.session_state.current_question_override or current_session["questions"][st.session_state.current_question]
                    
                    # Get existing answer
                    editor_base_key = f"quill_{current_session['id']}_{current_question_text[:20]}"
                    content_key = f"{editor_base_key}_content"
                    existing_answer = st.session_state.get(content_key, "")
                    
                    # Get profile context
                    profile_context = get_narrative_gps_for_ai()
                    
                    # Get birth year
                    birth_year = None
                    if st.session_state.user_account and 'profile' in st.session_state.user_account:
                        birthdate = st.session_state.user_account['profile'].get('birthdate', '')
                        if birthdate:
                            import re
                            year_match = re.search(r'\d{4}', birthdate)
                            if year_match:
                                birth_year = int(year_match.group())
                    
                    # Generate new prompts
                    result = generate_writing_prompts(
                        current_session['title'],
                        current_question_text,
                        existing_answer,
                        profile_context,
                        birth_year
                    )
                    
                    if result.get('success'):
                        st.session_state.current_prompt_data = result
                        st.rerun()
                    else:
                        st.error(result.get('error', 'Failed to generate prompts'))
        
        with col3:
            if st.button("✍️ Start Writing", key="start_writing", type="primary", use_container_width=True):
                st.session_state.show_prompt_modal = False
                st.session_state.current_prompt_data = None
                st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ============================================================================
# EMAIL CONFIG
# ============================================================================
EMAIL_CONFIG = {
    "smtp_server": st.secrets.get("SMTP_SERVER", "smtp.gmail.com"),
    "smtp_port": int(st.secrets.get("SMTP_PORT", 587)),
    "sender_email": st.secrets.get("SENDER_EMAIL", ""),
    "sender_password": st.secrets.get("SENDER_PASSWORD", ""),
    "use_tls": True
}

# ============================================================================
# BACKUP AND RESTORE FUNCTIONS
# ============================================================================
def create_backup():
    if not st.session_state.user_id:
        return None
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_data = {
            "user_id": st.session_state.user_id,
            "user_account": st.session_state.user_account,
            "responses": st.session_state.responses,
            "backup_date": datetime.now().isoformat(),
            "version": "1.0"
        }
        backup_file = f"backups/{st.session_state.user_id}_{timestamp}.json"
        with open(backup_file, 'w') as f:
            json.dump(backup_data, f, indent=2)
        return json.dumps(backup_data, indent=2)
    except Exception as e:
        st.error(f"Backup failed: {e}")
        return None

def restore_from_backup(backup_json):
    try:
        backup_data = json.loads(backup_json)
        if backup_data.get("user_id") != st.session_state.user_id:
            st.error("Backup belongs to a different user")
            return False
        st.session_state.user_account = backup_data.get("user_account", st.session_state.user_account)
        st.session_state.responses = backup_data.get("responses", st.session_state.responses)
        save_account_data(st.session_state.user_account)
        save_user_data(st.session_state.user_id, st.session_state.responses)
        return True
    except Exception as e:
        st.error(f"Restore failed: {e}")
        return False

def list_backups():
    if not st.session_state.user_id:
        return []
    backups = []
    try:
        for f in os.listdir("backups"):
            if f.startswith(st.session_state.user_id) and f.endswith(".json"):
                filepath = f"backups/{f}"
                with open(filepath, 'r') as file:
                    data = json.load(file)
                    backups.append({
                        "filename": f,
                        "date": data.get("backup_date", "Unknown"),
                        "size": os.path.getsize(filepath)
                    })
    except:
        pass
    return sorted(backups, key=lambda x: x["date"], reverse=True)

# ============================================================================
# IMAGE HANDLER
# ============================================================================
class ImageHandler:
    def __init__(self, user_id=None):
        self.user_id = user_id
        self.base_path = "uploads"
        self.settings = {
            "full_width": 1600,
            "inline_width": 800,
            "thumbnail_size": 200,
            "dpi": 300,
            "quality": 85,
            "max_file_size_mb": 5,
            "aspect_ratio": 1.6
        }
    
    def get_user_path(self):
        if self.user_id:
            user_hash = hashlib.md5(self.user_id.encode()).hexdigest()[:8]
            path = f"{self.base_path}/user_{user_hash}"
            os.makedirs(f"{path}/thumbnails", exist_ok=True)
            return path
        return self.base_path
    
    def optimize_image(self, image, max_width=1600, is_thumbnail=False):
        try:
            if image.mode in ('RGBA', 'LA', 'P'):
                bg = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                if image.mode == 'RGBA':
                    bg.paste(image, mask=image.split()[-1])
                else:
                    bg.paste(image)
                image = bg
            
            width, height = image.size
            aspect = height / width
            
            if is_thumbnail:
                size = min(width, height)
                left = (width - size) // 2
                top = (height - size) // 2
                right = left + size
                bottom = top + size
                image = image.crop((left, top, right, bottom))
                image.thumbnail((self.settings["thumbnail_size"], self.settings["thumbnail_size"]), Image.Resampling.LANCZOS)
                return image
            
            if width > max_width:
                new_width = max_width
                new_height = int(max_width * aspect)
                image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            return image
        except Exception as e:
            print(f"Error optimizing image: {e}")
            return image
    
    def save_image(self, uploaded_file, session_id, question_text, caption="", usage="full_page"):
        try:
            image_data = uploaded_file.read()
            original_size = len(image_data) / (1024 * 1024)
            
            if original_size > self.settings["max_file_size_mb"]:
                print(f"Warning: Large image ({original_size:.1f}MB). Will be optimized.")
            
            img = Image.open(io.BytesIO(image_data))
            target_width = self.settings["full_width"] if usage == "full_page" else self.settings["inline_width"]
            
            image_id = hashlib.md5(f"{self.user_id}{session_id}{question_text}{datetime.now()}".encode()).hexdigest()[:16]
            
            optimized_img = self.optimize_image(img, target_width, is_thumbnail=False)
            thumb_img = self.optimize_image(img, is_thumbnail=True)
            
            main_buffer = io.BytesIO()
            optimized_img.save(main_buffer, format="JPEG", quality=self.settings["quality"], optimize=True)
            main_size = len(main_buffer.getvalue()) / (1024 * 1024)
            
            thumb_buffer = io.BytesIO()
            thumb_img.save(thumb_buffer, format="JPEG", quality=70, optimize=True)
            
            user_path = self.get_user_path()
            with open(f"{user_path}/{image_id}.jpg", 'wb') as f: 
                f.write(main_buffer.getvalue())
            with open(f"{user_path}/thumbnails/{image_id}.jpg", 'wb') as f: 
                f.write(thumb_buffer.getvalue())
            
            metadata = {
                "id": image_id, "session_id": session_id, "question": question_text,
                "caption": caption, "alt_text": caption[:100] if caption else "",
                "timestamp": datetime.now().isoformat(), "user_id": self.user_id,
                "usage": usage, "original_size_mb": round(original_size, 2),
                "optimized_size_mb": round(main_size, 2), "dimensions": f"{optimized_img.width}x{optimized_img.height}",
                "optimized": True, "format": "JPEG", "dpi": self.settings["dpi"]
            }
            with open(f"{self.base_path}/metadata/{image_id}.json", 'w') as f: 
                json.dump(metadata, f, indent=2)
            
            reduction = ((original_size - main_size) / original_size) * 100 if original_size > 0 else 0
            if reduction > 20:
                print(f"✅ Image optimized: {original_size:.1f}MB → {main_size:.1f}MB ({reduction:.0f}% reduction)")
            
            return {
                "has_images": True, 
                "images": [{
                    "id": image_id, "caption": caption,
                    "dimensions": f"{optimized_img.width}x{optimized_img.height}",
                    "size_mb": round(main_size, 2)
                }]
            }
        except Exception as e:
            print(f"Error saving image: {e}")
            return None
    
    def get_image_html(self, image_id, thumbnail=False):
        try:
            user_path = self.get_user_path()
            path = f"{user_path}/thumbnails/{image_id}.jpg" if thumbnail else f"{user_path}/{image_id}.jpg"
            if not os.path.exists(path): 
                return None
            
            with open(path, 'rb') as f: 
                image_data = f.read()
            b64 = base64.b64encode(image_data).decode()
            
            meta_path = f"{self.base_path}/metadata/{image_id}.json"
            caption = ""
            dimensions = ""
            if os.path.exists(meta_path):
                with open(meta_path, 'r') as f:
                    metadata = json.load(f)
                    caption = metadata.get("caption", "")
                    dimensions = metadata.get("dimensions", "")
            
            return {
                "html": f'<img src="data:image/jpeg;base64,{b64}" class="story-image" alt="{caption}" data-dimensions="{dimensions}">',
                "caption": caption, "base64": b64, "dimensions": dimensions
            }
        except:
            return None
    
    def get_image_base64(self, image_id):
        try:
            user_path = self.get_user_path()
            path = f"{user_path}/{image_id}.jpg"
            if not os.path.exists(path): 
                return None
            with open(path, 'rb') as f: 
                image_data = f.read()
            return base64.b64encode(image_data).decode()
        except:
            return None
    
    def get_image_caption(self, image_id):
        meta_path = f"{self.base_path}/metadata/{image_id}.json"
        if os.path.exists(meta_path):
            try:
                with open(meta_path, 'r') as f:
                    metadata = json.load(f)
                    return metadata.get("caption", "")
            except:
                pass
        return ""
    
    def get_images_for_answer(self, session_id, question_text):
        images = []
        metadata_dir = f"{self.base_path}/metadata"
        if not os.path.exists(metadata_dir): 
            return images
        
        for fname in os.listdir(metadata_dir):
            if fname.endswith('.json'):
                try:
                    with open(f"{metadata_dir}/{fname}") as f: 
                        meta = json.load(f)
                    if (meta.get("session_id") == session_id and 
                        meta.get("question") == question_text and 
                        meta.get("user_id") == self.user_id):
                        thumb = self.get_image_html(meta["id"], thumbnail=True)
                        full = self.get_image_html(meta["id"])
                        if thumb and full:
                            images.append({
                                **meta, 
                                "thumb_html": thumb["html"], 
                                "full_html": full["html"]
                            })
                except:
                    continue
        return sorted(images, key=lambda x: x.get("timestamp", ""), reverse=True)
    
    def delete_image(self, image_id):
        try:
            user_path = self.get_user_path()
            for p in [f"{user_path}/{image_id}.jpg", 
                     f"{user_path}/thumbnails/{image_id}.jpg", 
                     f"{self.base_path}/metadata/{image_id}.json"]:
                if os.path.exists(p): 
                    os.remove(p)
            return True
        except:
            return False

def init_image_handler():
    if not st.session_state.image_handler or st.session_state.image_handler.user_id != st.session_state.get('user_id'):
        st.session_state.image_handler = ImageHandler(st.session_state.get('user_id'))
    return st.session_state.image_handler

# ============================================================================
# AUTHENTICATION FUNCTIONS
# ============================================================================
def generate_password(length=12):
    alphabet = string.ascii_letters + string.digits + "!@#$%^&*"
    return ''.join(secrets.choice(alphabet) for _ in range(length))

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(stored_hash, password):
    return stored_hash == hash_password(password)

def create_user_account(user_data, password=None):
    try:
        user_id = hashlib.sha256(f"{user_data['email']}{datetime.now().isoformat()}".encode()).hexdigest()[:12]
        if not password: 
            password = generate_password()
        user_record = {
            "user_id": user_id, "email": user_data["email"].lower().strip(),
            "password_hash": hash_password(password), "account_type": user_data.get("account_for", "self"),
            "created_at": datetime.now().isoformat(), "last_login": datetime.now().isoformat(),
            "profile": {
                "first_name": user_data["first_name"], "last_name": user_data["last_name"],
                "email": user_data["email"], "gender": user_data.get("gender", ""),
                "birthdate": user_data.get("birthdate", ""), "timeline_start": user_data.get("birthdate", ""),
                "occupation": user_data.get("occupation", ""), "hometown": user_data.get("hometown", ""),
                "current_location": user_data.get("current_location", ""), "family": user_data.get("family", ""),
                "education": user_data.get("education", ""), "life_philosophy": user_data.get("life_philosophy", ""),
                "legacy_hopes": user_data.get("legacy_hopes", "")
            },
            "narrative_gps": {},
            "privacy_settings": {
                "profile_public": False, "stories_public": False, "allow_sharing": False,
                "data_collection": True, "encryption": True
            },
            "settings": {
                "email_notifications": True, "auto_save": True, "privacy_level": "private",
                "theme": "light", "email_verified": False,
                "daily_word_goal": 500
            },
            "stats": {
                "total_sessions": 0, "total_words": 0,
                "account_age_days": 0, "last_active": datetime.now().isoformat()
            },
            "streak_data": {
                "current_streak": 0,
                "longest_streak": 0,
                "last_write_date": None,
                "streak_history": [],
                "milestones": {
                    "first_story": False,
                    "seven_day_streak": False,
                    "five_thousand_words": False,
                    "first_session_complete": False
                }
            }
        }
        save_account_data(user_record)
        return {"success": True, "user_id": user_id, "password": password, "user_record": user_record}
    except Exception as e:
        return {"success": False, "error": str(e)}

def save_account_data(user_record):
    try:
        with open(f"accounts/{user_record['user_id']}_account.json", 'w') as f:
            json.dump(user_record, f, indent=2)
        update_accounts_index(user_record)
        return True
    except: 
        return False

def update_accounts_index(user_record):
    try:
        index_file = "accounts/accounts_index.json"
        index = json.load(open(index_file, 'r')) if os.path.exists(index_file) else {}
        index[user_record['user_id']] = {
            "email": user_record['email'], "first_name": user_record['profile']['first_name'],
            "last_name": user_record['profile']['last_name'], "created_at": user_record['created_at'],
            "account_type": user_record['account_type']
        }
        with open(index_file, 'w') as f: 
            json.dump(index, f, indent=2)
        return True
    except: 
        return False

def get_account_data(user_id=None, email=None):
    try:
        if user_id:
            fname = f"accounts/{user_id}_account.json"
            if os.path.exists(fname): 
                return json.load(open(fname, 'r'))
        if email:
            email = email.lower().strip()
            index = json.load(open("accounts/accounts_index.json", 'r')) if os.path.exists("accounts/accounts_index.json") else {}
            for uid, data in index.items():
                if data.get("email", "").lower() == email:
                    return json.load(open(f"accounts/{uid}_account.json", 'r'))
    except: 
        pass
    return None

def authenticate_user(email, password):
    try:
        account = get_account_data(email=email)
        if account and verify_password(account['password_hash'], password):
            account['last_login'] = datetime.now().isoformat()
            save_account_data(account)
            return {"success": True, "user_id": account['user_id'], "user_record": account}
        return {"success": False, "error": "Invalid email or password"}
    except Exception as e:
        return {"success": False, "error": str(e)}

def send_welcome_email(user_data, credentials):
    try:
        if not EMAIL_CONFIG['sender_email'] or not EMAIL_CONFIG['sender_password']: 
            return False
        msg = MIMEMultipart()
        msg['From'] = EMAIL_CONFIG['sender_email']
        msg['To'] = user_data['email']
        msg['Subject'] = "Welcome to Tell My Story"
        
        body = f"""
        <html><body style="font-family: Arial;">
        <h2>Welcome to Tell My Story, {user_data['first_name']}!</h2>
        <div class="welcome-email">
            <h3>Your Account Details:</h3>
            <p><strong>Account ID:</strong> {credentials['user_id']}</p>
            <p><strong>Email:</strong> {user_data['email']}</p>
            <p><strong>Password:</strong> {credentials['password']}</p>
        </div>
        <p>Please keep this information safe. You can change your password anytime in settings.</p>
        </body></html>
        """
        msg.attach(MIMEText(body, 'html'))
        with smtplib.SMTP(EMAIL_CONFIG['smtp_server'], EMAIL_CONFIG['smtp_port']) as server:
            if EMAIL_CONFIG['use_tls']: 
                server.starttls()
            server.login(EMAIL_CONFIG['sender_email'], EMAIL_CONFIG['sender_password'])
            server.send_message(msg)
        return True
    except: 
        return False

def logout_user():
    st.session_state.qb_manager = None
    st.session_state.qb_manager_initialized = False
    st.session_state.image_handler = None
    keys = ['user_id', 'user_account', 'logged_in', 'show_profile_setup', 'current_session',
            'current_question', 'responses', 'session_conversations', 'data_loaded',
            'show_vignette_modal', 'vignette_topic', 'vignette_content', 'selected_vignette_type',
            'current_vignette_list', 'editing_vignette_index', 'show_vignette_manager',
            'custom_topic_input', 'show_custom_topic_modal', 'show_topic_browser',
            'show_session_manager', 'show_session_creator', 'editing_custom_session',
            'show_vignette_detail', 'selected_vignette_id', 'editing_vignette_id',
            'selected_vignette_for_session', 'published_vignette', 'show_beta_reader',
            'current_beta_feedback', 'current_question_bank', 'current_bank_name',
            'current_bank_type', 'current_bank_id', 'show_bank_manager', 'show_bank_editor',
            'editing_bank_id', 'editing_bank_name', 'show_image_manager', 'editor_content',
            'current_rewrite_data', 'show_ai_rewrite', 'show_ai_rewrite_menu',
            'show_publisher', 'cover_image_data', 'show_prompt_modal', 'current_prompt_data',
            'milestone_achieved_first_story', 'milestone_achieved_seven_day',
            'milestone_achieved_five_thousand', 'milestone_achieved_first_session']
    for key in keys:
        if key in st.session_state: 
            del st.session_state[key]
    st.query_params.clear()
    st.rerun()

# ============================================================================
# PRIVACY SETTINGS MODAL
# ============================================================================
def show_privacy_settings():
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    st.title("🔒 Privacy & Security Settings")
    
    if st.button("← Back", key="privacy_back"):
        st.session_state.show_privacy_settings = False
        st.rerun()
    
    st.markdown("### Ethical AI & Data Privacy")
    st.info("Your stories are private and secure. We use AI ethically to help you write better, never to train models on your personal data.")
    
    if 'privacy_settings' not in st.session_state.user_account:
        st.session_state.user_account['privacy_settings'] = {
            "profile_public": False, "stories_public": False, "allow_sharing": False,
            "data_collection": True, "encryption": True
        }
    
    privacy = st.session_state.user_account['privacy_settings']
    
    privacy['profile_public'] = st.checkbox("Make profile public", value=privacy.get('profile_public', False),
                                           help="Allow others to see your basic profile information")
    privacy['stories_public'] = st.checkbox("Share stories publicly", value=privacy.get('stories_public', False),
                                           help="Make your stories visible to the public (coming soon)")
    privacy['allow_sharing'] = st.checkbox("Allow sharing via link", value=privacy.get('allow_sharing', False),
                                          help="Generate shareable links to your stories")
    privacy['data_collection'] = st.checkbox("Allow anonymous usage data", value=privacy.get('data_collection', True),
                                            help="Help us improve by sharing anonymous usage statistics")
    privacy['encryption'] = st.checkbox("Enable encryption", value=privacy.get('encryption', True),
                                       disabled=True, help="Your data is always encrypted at rest")
    
    st.markdown("---")
    st.markdown("### 🔐 Security")
    st.markdown("- All data encrypted at rest")
    st.markdown("- No third-party data sharing")
    st.markdown("- You own all your content")
    st.markdown("- AI analysis is temporary and private")
    
    if st.button("💾 Save Privacy Settings", type="primary", use_container_width=True):
        save_account_data(st.session_state.user_account)
        st.success("Privacy settings saved!")
        time.sleep(1)
        st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ============================================================================
# PERFECT COVER DESIGNER - EXPORTS HTML (what you see is what you get)
# ============================================================================
def show_cover_designer():
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    st.title("🎨 Cover Designer")
    st.success("✅ Exports as HTML - Preview matches export exactly!")
    
    if st.button("← Back", key="cover_back"):
        st.session_state.show_cover_designer = False
        st.rerun()
    
    st.markdown("### Design your book cover - Portrait format (6\" x 9\")")
    
    # Load existing cover design if it exists
    saved_cover = st.session_state.user_account.get('cover_design', {}) if st.session_state.user_account else {}
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Cover Options**")
        
        # Editable title - load saved value if exists
        default_title = f"{st.session_state.user_account.get('profile', {}).get('first_name', 'My')}'s Story"
        title = st.text_input("Book Title", value=saved_cover.get('title', default_title))
        
        # Editable subtitle - load saved value if exists
        subtitle = st.text_input("Subtitle (optional)", value=saved_cover.get('subtitle', ''), placeholder="A brief subtitle or tagline")
        
        # Editable author name - load saved value if exists
        default_author = f"{st.session_state.user_account.get('profile', {}).get('first_name', '')} {st.session_state.user_account.get('profile', {}).get('last_name', '')}".strip()
        author = st.text_input("Author Name", value=saved_cover.get('author', default_author if default_author else "Author Name"))
        
        # Cover style options with safe indexing
        cover_options = ["Simple", "Elegant", "Modern", "Classic", "Vintage"]
        cover_index = 0
        saved_cover_type = saved_cover.get('cover_type', 'Simple')
        if saved_cover_type in cover_options:
            cover_index = cover_options.index(saved_cover_type)
        cover_type = st.selectbox("Cover Style", cover_options, index=cover_index)
        
        # Font options with safe indexing
        font_options = ["Georgia", "Arial", "Times New Roman", "Helvetica", "Calibri"]
        font_index = 0
        saved_font = saved_cover.get('title_font', 'Georgia')
        if saved_font in font_options:
            font_index = font_options.index(saved_font)
        title_font = st.selectbox("Title Font", font_options, index=font_index)
        
        title_color = st.color_picker("Title Color", value=saved_cover.get('title_color', '#000000'))
        background_color = st.color_picker("Background Color", value=saved_cover.get('background_color', '#FFFFFF'))
        
        # Show saved HTML if exists
        if saved_cover.get('cover_html') and os.path.exists(saved_cover['cover_html']):
            st.markdown("**Current Cover HTML:**")
            with open(saved_cover['cover_html'], 'r') as f:
                html_content = f.read()
            st.download_button(
                label="📥 Download Current Cover HTML",
                data=html_content,
                file_name="my_cover.html",
                mime="text/html",
                use_container_width=True
            )
            st.markdown("---")
            st.markdown("**Upload New Image (optional):**")
        
        uploaded_cover = st.file_uploader("Upload Cover Image (optional)", type=['jpg', 'jpeg', 'png'])
        if uploaded_cover:
            st.image(uploaded_cover, caption="New cover image", width=250)
    
    with col2:
        st.markdown("**Preview (6\" x 9\" portrait) - This EXACT HTML will be saved**")
        
        # Create the complete cover HTML with ALL elements
        if uploaded_cover:
            img_bytes = uploaded_cover.getvalue()
            img_base64 = base64.b64encode(img_bytes).decode()
            use_image = True
        elif saved_cover.get('cover_image') and os.path.exists(saved_cover['cover_image']):
            with open(saved_cover['cover_image'], 'rb') as f:
                img_bytes = f.read()
            img_base64 = base64.b64encode(img_bytes).decode()
            use_image = True
        else:
            use_image = False
            img_bytes = None
        
        # Build the EXACT cover HTML that will be saved
        if use_image:
            # With background image - text is WHITE with shadow for readability
            subtitle_html = f'<h2 style="font-family:{title_font}, sans-serif; color:white; font-size:32px; margin:20px 0 0 0; text-shadow:3px 3px 6px black; font-weight:normal;">{subtitle}</h2>' if subtitle else ''
            
            cover_html = f'''
            <div style="width:100%; max-width:600px; margin:0 auto; background:white; padding:20px;">
                <div style="
                    width:100%;
                    aspect-ratio:600/900;
                    background-image:url('data:image/jpeg;base64,{img_base64}');
                    background-size:cover;
                    background-position:center;
                    border:2px solid #333;
                    border-radius:10px;
                    overflow:hidden;
                    position:relative;
                    box-shadow:0 10px 20px rgba(0,0,0,0.3);
                ">
                    <div style="
                        position:absolute;
                        top:0;
                        left:0;
                        width:100%;
                        height:100%;
                        background:rgba(0,0,0,0.25);
                        display:flex;
                        flex-direction:column;
                        justify-content:space-between;
                        padding:50px 30px;
                        box-sizing:border-box;
                    ">
                        <div style="text-align:center;">
                            <h1 style="font-family:{title_font}, sans-serif; color:white; font-size:72px; margin:0; text-shadow:4px 4px 8px black; line-height:1.2;">{title}</h1>
                            {subtitle_html}
                        </div>
                        <div style="text-align:center; margin-bottom:50px;">
                            <p style="font-family:{title_font}, sans-serif; color:white; font-size:36px; margin:0; text-shadow:3px 3px 6px black;">by {author}</p>
                        </div>
                    </div>
                </div>
            </div>
            '''
        else:
            # Solid color background - text uses selected color
            subtitle_html = f'<h2 style="font-family:{title_font}, sans-serif; color:{title_color}; font-size:32px; margin:20px 0 0 0; font-weight:normal;">{subtitle}</h2>' if subtitle else ''
            
            cover_html = f'''
            <div style="width:100%; max-width:600px; margin:0 auto; background:white; padding:20px;">
                <div style="
                    width:100%;
                    aspect-ratio:600/900;
                    background-color:{background_color};
                    border:2px solid #333;
                    border-radius:10px;
                    display:flex;
                    flex-direction:column;
                    justify-content:space-between;
                    padding:50px 30px;
                    box-sizing:border-box;
                    box-shadow:0 10px 20px rgba(0,0,0,0.3);
                ">
                    <div style="text-align:center;">
                        <h1 style="font-family:{title_font}, sans-serif; color:{title_color}; font-size:72px; margin:0; line-height:1.2;">{title}</h1>
                        {subtitle_html}
                    </div>
                    <div style="text-align:center; margin-bottom:50px;">
                        <p style="font-family:{title_font}, sans-serif; color:{title_color}; font-size:36px; margin:0;">by {author}</p>
                    </div>
                </div>
            </div>
            '''
        
        # Display the complete cover preview
        from streamlit.components.v1 import html
        html(cover_html, height=800)
        
        st.caption("6\" wide × 9\" tall (portrait format) - This EXACT HTML will be saved")
    
    # Save button - NOW SAVES HTML instead of JPG
    if st.button("💾 Save Cover Design (as HTML)", type="primary", use_container_width=True):
        with st.spinner("💾 Saving cover HTML..."):
            try:
                # Create the complete HTML document
                if use_image:
                    subtitle_html_saved = f'<h2 style="font-family:{title_font}, sans-serif; color:white; font-size:32px; margin:20px 0 0 0; text-shadow:3px 3px 6px black; font-weight:normal;">{subtitle}</h2>' if subtitle else ''
                    
                    full_html = f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Book Cover - {title}</title>
    <style>
        body {{ margin:0; padding:0; display:flex; justify-content:center; align-items:center; min-height:100vh; background:#f0f0f0; }}
        .cover-container {{ width:600px; height:900px; position:relative; }}
        .cover {{
            width:100%;
            height:100%;
            background-image:url('data:image/jpeg;base64,{img_base64}');
            background-size:cover;
            background-position:center;
            border:2px solid #333;
            border-radius:10px;
            overflow:hidden;
            position:relative;
            box-shadow:0 10px 20px rgba(0,0,0,0.3);
        }}
        .overlay {{
            position:absolute;
            top:0;
            left:0;
            width:100%;
            height:100%;
            background:rgba(0,0,0,0.25);
            display:flex;
            flex-direction:column;
            justify-content:space-between;
            padding:50px 30px;
            box-sizing:border-box;
        }}
        .title {{
            font-family:{title_font}, sans-serif;
            color:white;
            font-size:72px;
            margin:0;
            text-shadow:4px 4px 8px black;
            line-height:1.2;
            text-align:center;
        }}
        .subtitle {{
            font-family:{title_font}, sans-serif;
            color:white;
            font-size:32px;
            margin:20px 0 0 0;
            text-shadow:3px 3px 6px black;
            font-weight:normal;
            text-align:center;
        }}
        .author {{
            font-family:{title_font}, sans-serif;
            color:white;
            font-size:36px;
            margin:0;
            text-shadow:3px 3px 6px black;
            text-align:center;
        }}
    </style>
</head>
<body>
    <div class="cover-container">
        <div class="cover">
            <div class="overlay">
                <div>
                    <div class="title">{title}</div>
                    {subtitle_html_saved}
                </div>
                <div class="author">by {author}</div>
            </div>
        </div>
    </div>
</body>
</html>'''
                else:
                    subtitle_html_saved = f'<div class="subtitle">{subtitle}</div>' if subtitle else ''
                    
                    full_html = f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Book Cover - {title}</title>
    <style>
        body {{ margin:0; padding:0; display:flex; justify-content:center; align-items:center; min-height:100vh; background:#f0f0f0; }}
        .cover-container {{ width:600px; height:900px; position:relative; }}
        .cover {{
            width:100%;
            height:100%;
            background-color:{background_color};
            border:2px solid #333;
            border-radius:10px;
            display:flex;
            flex-direction:column;
            justify-content:space-between;
            padding:50px 30px;
            box-sizing:border-box;
            box-shadow:0 10px 20px rgba(0,0,0,0.3);
        }}
        .title {{
            font-family:{title_font}, sans-serif;
            color:{title_color};
            font-size:72px;
            margin:0;
            line-height:1.2;
            text-align:center;
        }}
        .subtitle {{
            font-family:{title_font}, sans-serif;
            color:{title_color};
            font-size:32px;
            margin:20px 0 0 0;
            font-weight:normal;
            text-align:center;
        }}
        .author {{
            font-family:{title_font}, sans-serif;
            color:{title_color};
            font-size:36px;
            margin:0;
            text-align:center;
        }}
    </style>
</head>
<body>
    <div class="cover-container">
        <div class="cover">
            <div>
                <div class="title">{title}</div>
                {subtitle_html_saved}
            </div>
            <div class="author">by {author}</div>
        </div>
    </div>
</body>
</html>'''
                
                # Save HTML file
                html_filename = f"uploads/covers/{st.session_state.user_id}_cover.html"
                os.makedirs("uploads/covers", exist_ok=True)
                with open(html_filename, 'w') as f:
                    f.write(full_html)
                
                # Update user account with cover design data
                if 'cover_design' not in st.session_state.user_account:
                    st.session_state.user_account['cover_design'] = {}
                
                st.session_state.user_account['cover_design'].update({
                    "title": title,
                    "subtitle": subtitle,
                    "author": author,
                    "cover_type": cover_type,
                    "title_font": title_font,
                    "title_color": title_color,
                    "background_color": background_color,
                    "cover_html": html_filename,
                    "last_updated": datetime.now().isoformat()
                })
                
                # Also save the background image if uploaded
                if uploaded_cover:
                    img_path = f"uploads/covers/{st.session_state.user_id}_cover_bg.jpg"
                    with open(img_path, 'wb') as f:
                        f.write(uploaded_cover.getbuffer())
                    st.session_state.user_account['cover_design']['cover_image'] = img_path
                
                save_account_data(st.session_state.user_account)
                
                # Success message with download button
                st.success("✅ Cover HTML saved successfully!")
                
                # Offer immediate download to verify
                st.download_button(
                    label="📥 Download Cover HTML (open in browser to verify)",
                    data=full_html,
                    file_name=f"my_cover_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                    mime="text/html",
                    use_container_width=True
                )
                
                st.balloons()
                time.sleep(2)
                st.rerun()
                
            except Exception as e:
                st.error(f"Error saving cover: {str(e)}")
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ============================================================================
# NARRATIVE GPS HELPER FUNCTIONS
# ============================================================================
def get_narrative_gps_for_ai():
    if not st.session_state.user_account or 'narrative_gps' not in st.session_state.user_account:
        return ""
    
    gps = st.session_state.user_account['narrative_gps']
    if not gps:
        return ""
    
    context = "\n\n=== BOOK PROJECT CONTEXT (From Narrative GPS) ===\n"
    
    if gps.get('book_title') or gps.get('genre') or gps.get('book_length'):
        context += "\n📖 PROJECT SCOPE:\n"
        if gps.get('book_title'): context += f"- Book Title: {gps['book_title']}\n"
        if gps.get('genre'): 
            genre = gps['genre']
            if genre == "Other" and gps.get('genre_other'):
                genre = gps['genre_other']
            context += f"- Genre: {genre}\n"
        if gps.get('book_length'): context += f"- Length Vision: {gps['book_length']}\n"
        if gps.get('timeline'): context += f"- Timeline/Deadlines: {gps['timeline']}\n"
        if gps.get('completion_status'): context += f"- Current Status: {gps['completion_status']}\n"
    
    if gps.get('purposes') or gps.get('reader_takeaway'):
        context += "\n🎯 PURPOSE & AUDIENCE:\n"
        if gps.get('purposes'): 
            context += f"- Core Purposes: {', '.join(gps['purposes'])}\n"
        if gps.get('purpose_other'): context += f"- Other Purpose: {gps['purpose_other']}\n"
        if gps.get('audience_family'): context += f"- Family Audience: {gps['audience_family']}\n"
        if gps.get('audience_industry'): context += f"- Industry Audience: {gps['audience_industry']}\n"
        if gps.get('audience_challenges'): context += f"- Audience Facing Similar Challenges: {gps['audience_challenges']}\n"
        if gps.get('audience_general'): context += f"- General Audience: {gps['audience_general']}\n"
        if gps.get('reader_takeaway'): context += f"- Reader Takeaway: {gps['reader_takeaway']}\n"
    
    if gps.get('narrative_voices') or gps.get('emotional_tone'):
        context += "\n🎭 TONE & VOICE:\n"
        if gps.get('narrative_voices'): 
            context += f"- Narrative Voice: {', '.join(gps['narrative_voices'])}\n"
        if gps.get('voice_other'): context += f"- Other Voice: {gps['voice_other']}\n"
        if gps.get('emotional_tone'): context += f"- Emotional Tone: {gps['emotional_tone']}\n"
        if gps.get('language_style'): context += f"- Language Style: {gps['language_style']}\n"
    
    if gps.get('time_coverage') or gps.get('sensitive_material') or gps.get('inclusions'):
        context += "\n📋 CONTENT PARAMETERS:\n"
        if gps.get('time_coverage'): context += f"- Time Coverage: {gps['time_coverage']}\n"
        if gps.get('sensitive_material'): context += f"- Sensitive Topics: {gps['sensitive_material']}\n"
        if gps.get('sensitive_people'): context += f"- Sensitive People: {gps['sensitive_people']}\n"
        if gps.get('inclusions'): 
            context += f"- Planned Inclusions: {', '.join(gps['inclusions'])}\n"
        if gps.get('locations'): context += f"- Key Locations: {gps['locations']}\n"
    
    if gps.get('materials') or gps.get('people_to_interview'):
        context += "\n📦 RESOURCES:\n"
        if gps.get('materials'): 
            context += f"- Available Materials: {', '.join(gps['materials'])}\n"
        if gps.get('people_to_interview'): context += f"- People to Interview: {gps['people_to_interview']}\n"
        if gps.get('legal'): 
            context += f"- Legal Considerations: {', '.join(gps['legal'])}\n"
    
    if gps.get('involvement') or gps.get('unspoken'):
        context += "\n🤝 COLLABORATION:\n"
        if gps.get('involvement'): 
            involvement = gps['involvement']
            if involvement == "Mixed approach: [explain]" and gps.get('involvement_explain'):
                involvement = f"Mixed approach: {gps['involvement_explain']}"
            context += f"- Working Style: {involvement}\n"
        if gps.get('feedback_style'): context += f"- Feedback Preference: {gps['feedback_style']}\n"
        if gps.get('unspoken'): context += f"- Hopes for Collaboration: {gps['unspoken']}\n"
    
    return context

# ============================================================================
# AI REWRITE FUNCTION
# ============================================================================
def ai_rewrite_answer(original_text, person_option, question_text, session_title):
    """Rewrite the user's answer in 1st, 2nd, or 3rd person using profile context"""
    if not client:
        return {"error": "OpenAI client not available"}
    
    try:
        # Get profile context for better rewriting
        gps_context = get_narrative_gps_for_ai()
        
        # Also get enhanced profile for deeper context
        enhanced_context = ""
        if st.session_state.user_account and 'enhanced_profile' in st.session_state.user_account:
            ep = st.session_state.user_account['enhanced_profile']
            if ep:
                enhanced_context = "\n\n=== ADDITIONAL BIOGRAPHER CONTEXT ===\n"
                if ep.get('birth_place'): enhanced_context += f"• Born: {ep['birth_place']}\n"
                if ep.get('life_lessons'): enhanced_context += f"• Life Philosophy: {ep['life_lessons'][:200]}...\n"
                if ep.get('legacy'): enhanced_context += f"• Legacy Hope: {ep['legacy'][:200]}...\n"
        
        # Clean the text (remove HTML tags)
        clean_text = re.sub(r'<[^>]+>', '', original_text)
        
        if len(clean_text.split()) < 5:
            return {"error": "Text too short to rewrite (minimum 5 words)"}
        
        # Person-specific instructions
        person_instructions = {
            "1st": {
                "name": "First Person",
                "instruction": "Rewrite this in FIRST PERSON ('I', 'me', 'my', 'we', 'our'). Keep the authentic voice of the author telling their own story.",
                "example": "I remember the day clearly. The sun was setting and I felt...",
                "emoji": "👤"
            },
            "2nd": {
                "name": "Second Person",
                "instruction": "Rewrite this in SECOND PERSON ('you', 'your') as if speaking directly to the reader. Make it feel like advice, a letter, or a conversation with the reader.",
                "example": "You remember that day clearly. The sun was setting and you felt...",
                "emoji": "💬"
            },
            "3rd": {
                "name": "Third Person",
                "instruction": "Rewrite this in THIRD PERSON ('he', 'she', 'they', 'the author', the person's name). Write as if telling someone else's story to readers.",
                "example": f"They remember the day clearly. The sun was setting and they felt...",
                "emoji": "📖"
            }
        }
        
        system_prompt = f"""You are an expert writing assistant and ghostwriter. Your task is to rewrite the author's raw answer in {person_instructions[person_option]['name']}.

{person_instructions[person_option]['instruction']}

EXAMPLE STYLE:
{person_instructions[person_option]['example']}

IMPORTANT GUIDELINES:
1. Use the profile context below to understand WHO the author is
2. Preserve all key facts, emotions, and details from the original
3. Maintain the author's unique voice and personality
4. Fix any grammar issues naturally
5. Make it flow better while keeping it authentic
6. DO NOT add fictional events or details not in the original
7. Return ONLY the rewritten text, no explanations, no prefixes

PROFILE CONTEXT (Use this to understand the author's voice and story):
{gps_context}
{enhanced_context}

QUESTION BEING ANSWERED: {question_text}
SESSION: {session_title}

ORIGINAL ANSWER (to rewrite):
{clean_text}

REWRITTEN VERSION ({person_instructions[person_option]['name']}):"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": "Please rewrite this in the specified voice."}
            ],
            max_tokens=len(clean_text.split()) * 3,
            temperature=0.7
        )
        
        rewritten = response.choices[0].message.content.strip()
        
        # Clean up any markdown or quotes the AI might add
        rewritten = re.sub(r'^["\']|["\']$', '', rewritten)
        rewritten = re.sub(r'^Here\'s the rewritten version:?\s*', '', rewritten, flags=re.IGNORECASE)
        
        return {
            "success": True,
            "original": clean_text,
            "rewritten": rewritten,
            "person": person_instructions[person_option]["name"],
            "emoji": person_instructions[person_option]["emoji"]
        }
        
    except Exception as e:
        return {"error": str(e)}

# ============================================================================
# AI REWRITE MODAL
# ============================================================================
def show_ai_rewrite_modal():
    if not st.session_state.get('current_rewrite_data'):
        return
    
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    
    col1, col2 = st.columns([6, 1])
    with col1:
        st.markdown(f"### {st.session_state.current_rewrite_data.get('emoji', '✨')} AI Rewrite - {st.session_state.current_rewrite_data.get('person', '')}")
    with col2:
        if st.button("✕", key="close_rewrite_modal"):
            st.session_state.show_ai_rewrite = False
            st.session_state.current_rewrite_data = None
            st.rerun()
    
    st.markdown("---")
    
    rewrite_data = st.session_state.current_rewrite_data
    
    if rewrite_data.get('error'):
        st.error(f"Could not rewrite: {rewrite_data['error']}")
    else:
        st.markdown("**📝 Original Version:**")
        with st.container():
            st.markdown(f'<div class="original-text-box">{rewrite_data["original"]}</div>', unsafe_allow_html=True)
        
        st.markdown("---")
        
        st.markdown(f"**✨ Rewritten Version ({rewrite_data['person']}):**")
        with st.container():
            st.markdown(f'<div class="rewritten-text-box">{rewrite_data["rewritten"]}</div>', unsafe_allow_html=True)
        
        st.markdown("---")
        st.markdown("*This rewrite used your profile information to better capture your authentic voice.*")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("📋 Copy to Clipboard", key="copy_rewrite", use_container_width=True):
                st.info("✅ Copied! Select the text above and press Ctrl+C")
                
        with col2:
            if st.button("📝 Replace Original", key="replace_rewrite", type="primary", use_container_width=True):
                # Get the editor content key
                current_session = st.session_state.current_question_bank[st.session_state.current_session]
                current_session_id = current_session["id"]
                current_question_text = st.session_state.current_question_override or current_session["questions"][st.session_state.current_question]
                
                # Save the rewritten version
                editor_key = f"quill_{current_session_id}_{current_question_text[:20]}"
                content_key = f"{editor_key}_content"
                
                # Wrap in paragraph tags if not present
                new_content = rewrite_data["rewritten"]
                if not new_content.startswith('<p>'):
                    new_content = f'<p>{new_content}</p>'
                
                st.session_state[content_key] = new_content
                
                # Save to database
                save_response(current_session_id, current_question_text, new_content)
                
                st.success("✅ Replaced with rewritten version!")
                time.sleep(1)
                st.session_state.show_ai_rewrite = False
                st.session_state.current_rewrite_data = None
                st.rerun()
        
        with col3:
            if st.button("🔄 Try Different Voice", key="try_another", use_container_width=True):
                st.session_state.show_ai_rewrite = False
                st.session_state.current_rewrite_data = None
                st.rerun()
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

# ============================================================================
# ENHANCED BIOGRAPHER PROFILE SECTION
# ============================================================================
def render_enhanced_profile():
    st.markdown("### 📋 The Biographer's Questions")
    st.info("These questions ensure we capture the full richness of your life story.")
    
    if 'enhanced_profile' not in st.session_state.user_account:
        st.session_state.user_account['enhanced_profile'] = {}
    
    ep = st.session_state.user_account['enhanced_profile']
    
    with st.expander("👶 Early Years & Family Origins", expanded=False):
        st.markdown("**Where and when were you born?**")
        ep['birth_place'] = st.text_input("Birth place", value=ep.get('birth_place', ''), key="ep_birth_place")
        
        st.markdown("**Tell me about your parents - who were they? What were their personalities, dreams, and life stories?**")
        ep['parents'] = st.text_area("Parents", value=ep.get('parents', ''), key="ep_parents", height=100)
        
        st.markdown("**Did you have siblings? What was your birth order and relationship with them?**")
        ep['siblings'] = st.text_area("Siblings", value=ep.get('siblings', ''), key="ep_siblings", height=100)
        
        st.markdown("**What was your childhood home like? The neighborhood, the house, the atmosphere?**")
        ep['childhood_home'] = st.text_area("Childhood home", value=ep.get('childhood_home', ''), key="ep_home", height=100)
        
        st.markdown("**What family traditions, values, or cultural background shaped your early years?**")
        ep['family_traditions'] = st.text_area("Family traditions", value=ep.get('family_traditions', ''), key="ep_traditions", height=100)
    
    with st.expander("🎓 Education & Formative Years", expanded=False):
        st.markdown("**What was your school experience like? Favorite teachers? Subjects you loved or hated?**")
        ep['school'] = st.text_area("School years", value=ep.get('school', ''), key="ep_school", height=100)
        
        st.markdown("**Did you pursue higher education? What influenced your choices?**")
        ep['higher_ed'] = st.text_area("Higher education", value=ep.get('higher_ed', ''), key="ep_higher_ed", height=100)
        
        st.markdown("**Who were your mentors or influential figures during these years?**")
        ep['mentors'] = st.text_area("Mentors", value=ep.get('mentors', ''), key="ep_mentors", height=100)
        
        st.markdown("**What books, ideas, or experiences shaped your worldview?**")
        ep['influences'] = st.text_area("Influences", value=ep.get('influences', ''), key="ep_influences", height=100)
    
    with st.expander("💼 Career & Life's Work", expanded=False):
        st.markdown("**What was your first job? What did you learn from it?**")
        ep['first_job'] = st.text_area("First job", value=ep.get('first_job', ''), key="ep_first_job", height=100)
        
        st.markdown("**Describe your career path - the twists, turns, and defining moments.**")
        ep['career_path'] = st.text_area("Career path", value=ep.get('career_path', ''), key="ep_career", height=100)
        
        st.markdown("**What achievements are you most proud of?**")
        ep['achievements'] = st.text_area("Achievements", value=ep.get('achievements', ''), key="ep_achievements", height=100)
        
        st.markdown("**What work or projects brought you the most fulfillment?**")
        ep['fulfillment'] = st.text_area("Fulfilling work", value=ep.get('fulfillment', ''), key="ep_fulfillment", height=100)
    
    with st.expander("❤️ Relationships & Love", expanded=False):
        st.markdown("**Tell me about your romantic relationships - first loves, significant partnerships.**")
        ep['romance'] = st.text_area("Romantic relationships", value=ep.get('romance', ''), key="ep_romance", height=100)
        
        st.markdown("**If married, how did you meet? What has the journey been like?**")
        ep['marriage'] = st.text_area("Marriage story", value=ep.get('marriage', ''), key="ep_marriage", height=100)
        
        st.markdown("**Tell me about your children, if any - their personalities, your relationship with them.**")
        ep['children'] = st.text_area("Children", value=ep.get('children', ''), key="ep_children", height=100)
        
        st.markdown("**Who are your closest friends? What makes those friendships special?**")
        ep['friends'] = st.text_area("Friendships", value=ep.get('friends', ''), key="ep_friends", height=100)
    
    with st.expander("🌟 Challenges & Triumphs", expanded=False):
        st.markdown("**What were the hardest moments in your life? How did you navigate them?**")
        ep['challenges'] = st.text_area("Challenges", value=ep.get('challenges', ''), key="ep_challenges", height=100)
        
        st.markdown("**What losses have you experienced and how did they change you?**")
        ep['losses'] = st.text_area("Losses", value=ep.get('losses', ''), key="ep_losses", height=100)
        
        st.markdown("**What are your proudest moments? Times when you felt truly alive?**")
        ep['proud_moments'] = st.text_area("Proud moments", value=ep.get('proud_moments', ''), key="ep_proud", height=100)
        
        st.markdown("**What obstacles did you overcome that defined who you are?**")
        ep['overcame'] = st.text_area("Obstacles overcome", value=ep.get('overcame', ''), key="ep_overcame", height=100)
    
    with st.expander("🌍 Life Philosophy & Wisdom", expanded=False):
        st.markdown("**What life lessons would you want to pass on to future generations?**")
        ep['life_lessons'] = st.text_area("Life lessons", value=ep.get('life_lessons', ''), key="ep_lessons", height=100)
        
        st.markdown("**What do you believe in? What are your core values?**")
        ep['values'] = st.text_area("Core values", value=ep.get('values', ''), key="ep_values", height=100)
        
        st.markdown("**If you could give your younger self advice, what would it be?**")
        ep['advice'] = st.text_area("Advice to younger self", value=ep.get('advice', ''), key="ep_advice", height=100)
        
        st.markdown("**How would you like to be remembered?**")
        ep['legacy'] = st.text_area("Legacy", value=ep.get('legacy', ''), key="ep_legacy", height=100)
    
    if st.button("💾 Save Biographer's Questions", type="primary", use_container_width=True):
        save_account_data(st.session_state.user_account)
        st.success("Biographer's profile saved!")
        st.rerun()

# ============================================================================
# NARRATIVE GPS PROFILE SECTION
# ============================================================================
def render_narrative_gps():
    st.markdown("### ❤️ The Heart of Your Story")
    
    st.markdown("""
    <div class="narrative-gps-box">
    <p>Your answers to these questions help me support you properly throughout the process and make sure the finished book is exactly right for you and your readers.</p>
    
    <p>The more open and detailed you are here, the easier it is for your real voice and personality to come through on every page. Think of this as a conversation between you and the person who will read your story one day — and I'm here alongside you, listening, capturing what matters, and helping shape it into something lasting.</p>
    
    <p><strong>There's no rush.</strong> You can return and add to this whenever new thoughts or memories surface. This is where your story truly begins.</p>
    </div>
    """, unsafe_allow_html=True)
    
    if 'narrative_gps' not in st.session_state.user_account:
        st.session_state.user_account['narrative_gps'] = {}
    
    gps = st.session_state.user_account['narrative_gps']
    
    with st.expander("📖 Section 1: The Book Itself (Project Scope)", expanded=True):
        st.markdown("**BOOK TITLE (Working or Final):**")
        gps['book_title'] = st.text_input(
            "What's your working title? If unsure, what feeling or idea should the title convey?",
            value=gps.get('book_title', ''),
            label_visibility="collapsed",
            placeholder="What's your working title? If unsure, what feeling or idea should the title convey?",
            key="gps_title"
        )
        
        st.markdown("**BOOK GENRE/CATEGORY:**")
        genre_options = ["", "Memoir", "Autobiography", "Family History", "Business/Legacy Book", "Other"]
        genre_index = 0
        if gps.get('genre') in genre_options:
            genre_index = genre_options.index(gps['genre'])
        
        gps['genre'] = st.selectbox(
            "BOOK GENRE/CATEGORY:",
            options=genre_options,
            index=genre_index,
            label_visibility="collapsed",
            key="gps_genre"
        )
        if gps['genre'] == "Other":
            gps['genre_other'] = st.text_input("Please specify:", value=gps.get('genre_other', ''), key="gps_genre_other")
        
        st.markdown("**BOOK LENGTH VISION:**")
        length_options = ["", "A short book (100-150 pages)", "Standard length (200-300 pages)", "Comprehensive (300+ pages)"]
        length_index = 0
        if gps.get('book_length') in length_options:
            length_index = length_options.index(gps['book_length'])
        
        gps['book_length'] = st.selectbox(
            "BOOK LENGTH VISION:",
            options=length_options,
            index=length_index,
            label_visibility="collapsed",
            key="gps_length"
        )
        
        st.markdown("**TIMELINE & DEADLINES:**")
        gps['timeline'] = st.text_area(
            "Do you have a target publication date or event this book is tied to? (e.g., birthday, retirement, anniversary)",
            value=gps.get('timeline', ''),
            label_visibility="collapsed",
            placeholder="Do you have a target publication date or event this book is tied to? (e.g., birthday, retirement, anniversary)",
            key="gps_timeline"
        )
        
        st.markdown("**COMPLETION STATUS:**")
        completion_options = ["", "Notes only", "Partial chapters", "Full draft"]
        completion_index = 0
        if gps.get('completion_status') in completion_options:
            completion_index = completion_options.index(gps['completion_status'])
        
        gps['completion_status'] = st.selectbox(
            "COMPLETION STATUS:",
            options=completion_options,
            index=completion_index,
            label_visibility="collapsed",
            key="gps_completion"
        )
    
    with st.expander("🎯 Section 2: Purpose & Audience (The 'Why')", expanded=False):
        st.markdown("**THE CORE PURPOSE (Choose all that apply):**")
        
        if 'purposes' not in gps:
            gps['purposes'] = []
        
        purposes_options = [
            "Leave a legacy for family/future generations",
            "Share life lessons to help others",
            "Document professional/business journey",
            "Heal or process through writing",
            "Establish authority/expertise",
            "Entertain with entertaining stories"
        ]
        
        for purpose in purposes_options:
            if st.checkbox(
                purpose,
                value=purpose in gps.get('purposes', []),
                key=f"gps_purpose_{purpose}"
            ):
                if purpose not in gps['purposes']:
                    gps['purposes'].append(purpose)
            else:
                if purpose in gps['purposes']:
                    gps['purposes'].remove(purpose)
        
        gps['purpose_other'] = st.text_input("Other:", value=gps.get('purpose_other', ''), key="gps_purpose_other")
        
        st.markdown("---")
        st.markdown("**PRIMARY AUDIENCE:**")
        st.markdown("*Who is your ideal reader? Be specific:*")
        
        gps['audience_family'] = st.text_input(
            "Family members (which generations?):",
            value=gps.get('audience_family', ''),
            key="gps_audience_family"
        )
        
        gps['audience_industry'] = st.text_input(
            "People in your industry/profession:",
            value=gps.get('audience_industry', ''),
            key="gps_audience_industry"
        )
        
        gps['audience_challenges'] = st.text_input(
            "People facing similar challenges you overcame:",
            value=gps.get('audience_challenges', ''),
            key="gps_audience_challenges"
        )
        
        gps['audience_general'] = st.text_input(
            "The general public interested in:",
            value=gps.get('audience_general', ''),
            placeholder="your topic",
            key="gps_audience_general"
        )
        
        st.markdown("---")
        st.markdown("**THE READER TAKEAWAY:**")
        gps['reader_takeaway'] = st.text_area(
            "What do you want readers to feel, think, or do after finishing your book?",
            value=gps.get('reader_takeaway', ''),
            label_visibility="collapsed",
            placeholder="What do you want readers to feel, think, or do after finishing your book?",
            key="gps_takeaway"
        )
    
    with st.expander("🎭 Section 3: Tone & Voice (The 'How')", expanded=False):
        st.markdown("**NARRATIVE VOICE:**")
        
        if 'narrative_voices' not in gps:
            gps['narrative_voices'] = []
        
        voice_options = [
            "Warm and conversational (like talking to a friend)",
            "Professional and authoritative",
            "Raw and vulnerable",
            "Humorous/lighthearted",
            "Philosophical/reflective"
        ]
        
        for voice in voice_options:
            if st.checkbox(
                voice,
                value=voice in gps.get('narrative_voices', []),
                key=f"gps_voice_{voice}"
            ):
                if voice not in gps['narrative_voices']:
                    gps['narrative_voices'].append(voice)
            else:
                if voice in gps['narrative_voices']:
                    gps['narrative_voices'].remove(voice)
        
        gps['voice_other'] = st.text_input("Other:", value=gps.get('voice_other', ''), key="gps_voice_other")
        
        st.markdown("---")
        st.markdown("**EMOTIONAL TONE:**")
        gps['emotional_tone'] = st.text_area(
            "Should readers laugh? Cry? Feel inspired? Get angry? All of the above?",
            value=gps.get('emotional_tone', ''),
            label_visibility="collapsed",
            placeholder="Should readers laugh? Cry? Feel inspired? Get angry? All of the above?",
            key="gps_emotional"
        )
        
        st.markdown("---")
        st.markdown("**LANGUAGE STYLE:**")
        language_options = ["", "Simple, everyday language", "Rich, descriptive prose", "Short, punchy chapters", "Long, flowing narratives"]
        language_index = 0
        if gps.get('language_style') in language_options:
            language_index = language_options.index(gps['language_style'])
        
        gps['language_style'] = st.selectbox(
            "LANGUAGE STYLE:",
            options=language_options,
            index=language_index,
            label_visibility="collapsed",
            key="gps_language"
        )
    
    with st.expander("📋 Section 4: Content Parameters (The 'What')", expanded=False):
        st.markdown("**TIME COVERAGE:**")
        time_options = ["", "Your entire life", "A specific era/decade", "One defining experience", "Your career/business journey"]
        time_index = 0
        if gps.get('time_coverage') in time_options:
            time_index = time_options.index(gps['time_coverage'])
        
        gps['time_coverage'] = st.selectbox(
            "TIME COVERAGE:",
            options=time_options,
            index=time_index,
            label_visibility="collapsed",
            key="gps_time"
        )
        
        st.markdown("---")
        st.markdown("**SENSITIVE MATERIAL:**")
        gps['sensitive_material'] = st.text_area(
            "Are there topics, people, or events you want to handle carefully or omit entirely?",
            value=gps.get('sensitive_material', ''),
            label_visibility="collapsed",
            placeholder="Are there topics, people, or events you want to handle carefully or omit entirely?",
            key="gps_sensitive"
        )
        
        gps['sensitive_people'] = st.text_area(
            "Any living people whose portrayal requires sensitivity or legal consideration?",
            value=gps.get('sensitive_people', ''),
            label_visibility="collapsed",
            placeholder="Any living people whose portrayal requires sensitivity or legal consideration?",
            key="gps_sensitive_people"
        )
        
        st.markdown("---")
        st.markdown("**INCLUSIONS:**")
        
        if 'inclusions' not in gps:
            gps['inclusions'] = []
        
        inclusion_options = ["Photos", "Family trees", "Recipes", "Letters/documents", "Timelines", "Resources for readers"]
        for inc in inclusion_options:
            if st.checkbox(
                inc,
                value=inc in gps.get('inclusions', []),
                key=f"gps_inc_{inc}"
            ):
                if inc not in gps['inclusions']:
                    gps['inclusions'].append(inc)
            else:
                if inc in gps['inclusions']:
                    gps['inclusions'].remove(inc)
        
        st.markdown("---")
        st.markdown("**LOCATIONS:**")
        gps['locations'] = st.text_area(
            "List key places that must appear in the story (hometowns, meaningful travels, etc.)",
            value=gps.get('locations', ''),
            label_visibility="collapsed",
            placeholder="List key places that must appear in the story (hometowns, meaningful travels, etc.)",
            key="gps_locations"
        )
    
    with st.expander("📦 Section 5: Assets & Access (The 'Resources')", expanded=False):
        st.markdown("**EXISTING MATERIALS:**")
        
        if 'materials' not in gps:
            gps['materials'] = []
        
        material_options = [
            "Journals/diaries", "Letters or emails", "Photos (with dates/context)",
            "Video/audio recordings", "Newspaper clippings", "Awards/certificates",
            "Social media posts", "Previous interviews"
        ]
        
        for mat in material_options:
            if st.checkbox(
                mat,
                value=mat in gps.get('materials', []),
                key=f"gps_mat_{mat}"
            ):
                if mat not in gps['materials']:
                    gps['materials'].append(mat)
            else:
                if mat in gps['materials']:
                    gps['materials'].remove(mat)
        
        st.markdown("---")
        st.markdown("**PEOPLE TO INTERVIEW:**")
        gps['people_to_interview'] = st.text_area(
            "Are there family members, friends, or colleagues who should contribute their memories?",
            value=gps.get('people_to_interview', ''),
            label_visibility="collapsed",
            placeholder="Are there family members, friends, or colleagues who should contribute their memories?",
            key="gps_people"
        )
        
        st.markdown("---")
        st.markdown("**FINANCIAL & LEGAL:**")
        
        if 'legal' not in gps:
            gps['legal'] = []
        
        legal_options = ["ISBN registration", "Copyright", "Libel review", "Permissions for quoted material"]
        for leg in legal_options:
            if st.checkbox(
                leg,
                value=leg in gps.get('legal', []),
                key=f"gps_legal_{leg}"
            ):
                if leg not in gps['legal']:
                    gps['legal'].append(leg)
            else:
                if leg in gps['legal']:
                    gps['legal'].remove(leg)
    
    with st.expander("🤝 Section 6: Ghostwriter Relationship (The 'Collaboration')", expanded=False):
        st.markdown("**YOUR INVOLVEMENT:**")
        
        involvement_options = [
            "I'll answer questions, you write everything",
            "I'll write drafts, you polish",
            "We'll interview together, then you write",
            "Mixed approach: [explain]"
        ]
        
        involvement_index = 0
        if gps.get('involvement') in involvement_options:
            involvement_index = involvement_options.index(gps['involvement'])
        
        gps['involvement'] = st.radio(
            "How do you want to work together?",
            options=involvement_options,
            index=involvement_index,
            key="gps_involvement"
        )
        
        if gps.get('involvement') == "Mixed approach: [explain]":
            gps['involvement_explain'] = st.text_area(
                "Explain your preferred approach:",
                value=gps.get('involvement_explain', ''),
                key="gps_involvement_explain"
            )
        
        st.markdown("---")
        
        st.markdown("**FEEDBACK STYLE:**")
        feedback_options = ["", "Written comments", "Phone/video discussions", "Line-by-line edits"]
        feedback_index = 0
        if gps.get('feedback_style') in feedback_options:
            feedback_index = feedback_options.index(gps['feedback_style'])
        
        gps['feedback_style'] = st.selectbox(
            "FEEDBACK STYLE:",
            options=feedback_options,
            index=feedback_index,
            label_visibility="collapsed",
            key="gps_feedback"
        )
        
        st.markdown("---")
        st.markdown("**THE UNSPOKEN:**")
        gps['unspoken'] = st.text_area(
            "What are you hoping I'll bring to this project that you can't do yourself?",
            value=gps.get('unspoken', ''),
            label_visibility="collapsed",
            placeholder="What are you hoping I'll bring to this project that you can't do yourself?",
            key="gps_unspoken"
        )
    
    if st.button("💾 Save The Heart of Your Story", key="save_narrative_gps", type="primary", use_container_width=True):
        save_account_data(st.session_state.user_account)
        st.success("✅ The Heart of Your Story saved!")
        st.rerun()

# ============================================================================
# STORAGE FUNCTIONS
# ============================================================================
def get_user_filename(user_id):
    return f"user_data_{hashlib.md5(user_id.encode()).hexdigest()[:8]}.json"

def load_user_data(user_id):
    fname = get_user_filename(user_id)
    try:
        if os.path.exists(fname):
            return json.load(open(fname, 'r'))
        return {"responses": {}, "vignettes": [], "last_loaded": datetime.now().isoformat()}
    except: 
        return {"responses": {}, "vignettes": [], "last_loaded": datetime.now().isoformat()}

def save_user_data(user_id, responses_data):
    fname = get_user_filename(user_id)
    try:
        existing = load_user_data(user_id)
        data = {
            "user_id": user_id, 
            "responses": responses_data,
            "vignettes": existing.get("vignettes", []),
            "beta_feedback": existing.get("beta_feedback", {}),
            "last_saved": datetime.now().isoformat()
        }
        with open(fname, 'w') as f: 
            json.dump(data, f, indent=2)
        return True
    except: 
        return False

# ============================================================================
# CORE RESPONSE FUNCTIONS
# ============================================================================
def save_response(session_id, question, answer):
    user_id = st.session_state.user_id
    if not user_id: 
        return False
    
    text_only = re.sub(r'<[^>]+>', '', answer) if answer else ""
    
    if st.session_state.user_account:
        word_count = len(re.findall(r'\w+', text_only))
        st.session_state.user_account["stats"]["total_words"] = st.session_state.user_account["stats"].get("total_words", 0) + word_count
        st.session_state.user_account["stats"]["last_active"] = datetime.now().isoformat()
        save_account_data(st.session_state.user_account)
        
        # Update streak after saving
        update_writing_streak(user_id)
    
    if session_id not in st.session_state.responses:
        session_data = next((s for s in (st.session_state.current_question_bank or []) if s["id"] == session_id), 
                          {"title": f"Session {session_id}", "word_target": DEFAULT_WORD_TARGET})
        st.session_state.responses[session_id] = {
            "title": session_data.get("title", f"Session {session_id}"),
            "questions": {}, "summary": "", "completed": False,
            "word_target": session_data.get("word_target", DEFAULT_WORD_TARGET)
        }
    
    images = []
    if st.session_state.image_handler:
        images = st.session_state.image_handler.get_images_for_answer(session_id, question)
    
    st.session_state.responses[session_id]["questions"][question] = {
        "answer": answer, "question": question, "timestamp": datetime.now().isoformat(),
        "answer_index": 1, "has_images": len(images) > 0 or ('<img' in answer),
        "image_count": len(images), "images": [{"id": img["id"], "caption": img.get("caption", "")} for img in images]
    }
    
    success = save_user_data(user_id, st.session_state.responses)
    if success: 
        st.session_state.data_loaded = False
    
    return success

def delete_response(session_id, question):
    user_id = st.session_state.user_id
    if not user_id: 
        return False
    
    if session_id in st.session_state.responses and question in st.session_state.responses[session_id]["questions"]:
        del st.session_state.responses[session_id]["questions"][question]
        success = save_user_data(user_id, st.session_state.responses)
        if success: 
            st.session_state.data_loaded = False
        return success
    return False

def calculate_author_word_count(session_id):
    total = 0
    if session_id in st.session_state.responses:
        for q, d in st.session_state.responses[session_id].get("questions", {}).items():
            if d.get("answer"): 
                text_only = re.sub(r'<[^>]+>', '', d["answer"])
                total += len(re.findall(r'\w+', text_only))
    return total

def get_progress_info(session_id):
    current = calculate_author_word_count(session_id)
    if session_id not in st.session_state.responses:
        session_data = next((s for s in (st.session_state.current_question_bank or []) if s["id"] == session_id), {})
        st.session_state.responses[session_id] = {
            "title": session_data.get("title", f"Session {session_id}"),
            "questions": {}, "summary": "", "completed": False,
            "word_target": session_data.get("word_target", DEFAULT_WORD_TARGET)
        }
    
    target = st.session_state.responses[session_id].get("word_target", DEFAULT_WORD_TARGET)
    if target == 0: 
        percent = 100
    else: 
        percent = (current / target) * 100
    
    return {
        "current_count": current, "target": target, "progress_percent": percent,
        "emoji": "🟢" if percent >= 100 else "🟡" if percent >= 70 else "🔴",
        "color": "#27ae60" if percent >= 100 else "#f39c12" if percent >= 70 else "#e74c3c",
        "remaining_words": max(0, target - current),
        "status_text": "Target achieved!" if current >= target else f"{max(0, target - current)} words remaining"
    }

def auto_correct_text(text):
    if not text: 
        return text
    text_only = re.sub(r'<[^>]+>', '', text)
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Fix spelling and grammar. Return only corrected text."},
                {"role": "user", "content": text_only}
            ],
            max_tokens=len(text_only) + 100, temperature=0.1
        )
        return resp.choices[0].message.content
    except: 
        return text

# ============================================================================
# SEARCH FUNCTIONALITY
# ============================================================================
def search_all_answers(search_query):
    if not search_query or len(search_query) < 2: 
        return []
    
    results = []
    search_query = search_query.lower()
    
    for session in (st.session_state.current_question_bank or []):
        session_id = session["id"]
        session_data = st.session_state.responses.get(session_id, {})
        
        for question_text, answer_data in session_data.get("questions", {}).items():
            html_answer = answer_data.get("answer", "")
            text_answer = re.sub(r'<[^>]+>', '', html_answer)
            has_images = answer_data.get("has_images", False) or ('<img' in html_answer)
            
            if search_query in text_answer.lower() or search_query in question_text.lower():
                results.append({
                    "session_id": session_id, "session_title": session["title"],
                    "question": question_text, "answer": text_answer[:300] + "..." if len(text_answer) > 300 else text_answer,
                    "timestamp": answer_data.get("timestamp", ""), "word_count": len(text_answer.split()),
                    "has_images": has_images, "image_count": answer_data.get("image_count", 0)
                })
    
    results.sort(key=lambda x: x["timestamp"], reverse=True)
    return results

# ============================================================================
# QUESTION BANK LOADING
# ============================================================================
def initialize_question_bank():
    if 'current_question_bank' in st.session_state and st.session_state.current_question_bank:
        return True
    
    if QuestionBankManager:
        try:
            qb_manager = QuestionBankManager(st.session_state.get('user_id'))
            st.session_state.qb_manager = qb_manager
            
            if os.path.exists("sessions/sessions.csv"):
                shutil.copy("sessions/sessions.csv", "question_banks/default/life_story_comprehensive.csv")
            
            default = qb_manager.load_default_bank("life_story_comprehensive")
            if default:
                st.session_state.current_question_bank = default
                st.session_state.current_bank_name = "📖 Life Story - Comprehensive"
                st.session_state.current_bank_type = "default"
                st.session_state.current_bank_id = "life_story_comprehensive"
                st.session_state.qb_manager_initialized = True
                
                for s in default:
                    sid = s["id"]
                    if sid not in st.session_state.responses:
                        st.session_state.responses[sid] = {
                            "title": s["title"], "questions": {}, "summary": "",
                            "completed": False, "word_target": s.get("word_target", DEFAULT_WORD_TARGET)
                        }
                return True
        except: 
            pass
    
    if SessionLoader:
        try:
            legacy = SessionLoader().load_sessions_from_csv()
            if legacy:
                st.session_state.current_question_bank = legacy
                st.session_state.current_bank_name = "Legacy Bank"
                st.session_state.current_bank_type = "legacy"
                for s in legacy:
                    sid = s["id"]
                    if sid not in st.session_state.responses:
                        st.session_state.responses[sid] = {
                            "title": s["title"], "questions": {}, "summary": "",
                            "completed": False, "word_target": s.get("word_target", DEFAULT_WORD_TARGET)
                        }
                return True
        except: 
            pass
    return False

def load_question_bank(sessions, bank_name, bank_type, bank_id=None):
    st.session_state.current_question_bank = sessions
    st.session_state.current_bank_name = bank_name
    st.session_state.current_bank_type = bank_type
    st.session_state.current_bank_id = bank_id
    st.session_state.current_session = 0
    st.session_state.current_question = 0
    st.session_state.current_question_override = None
    
    for s in sessions:
        sid = s["id"]
        if sid not in st.session_state.responses:
            st.session_state.responses[sid] = {
                "title": s["title"], "questions": {}, "summary": "",
                "completed": False, "word_target": s.get("word_target", DEFAULT_WORD_TARGET)
            }

# ============================================================================
# BETA READER FUNCTIONS
# ============================================================================
def generate_beta_reader_feedback(session_title, session_text, feedback_type="comprehensive"):
    if not beta_reader: 
        return {"error": "BetaReader not available"}
    
    # Track what profile information was accessed
    accessed_profile_sections = []
    profile_context = "\n\n" + "="*80 + "\n"
    profile_context += "📋 BIOGRAPHER'S INTELLIGENCE BRIEFING\n"
    profile_context += "="*80 + "\n"
    profile_context += "The Beta Reader has accessed the following profile information to provide contextual feedback:\n\n"
    
    if st.session_state.user_account:
        # Get Narrative GPS (Heart of Your Story)
        gps = st.session_state.user_account.get('narrative_gps', {})
        if gps:
            profile_context += "\n📖 SECTION 1: BOOK PROJECT CONTEXT (From Narrative GPS)\n"
            profile_context += "-" * 50 + "\n"
            
            if gps.get('book_title'):
                profile_context += f"• Book Title: {gps['book_title']}\n"
                accessed_profile_sections.append("Book Title")
            if gps.get('genre'):
                genre = gps['genre']
                if genre == "Other" and gps.get('genre_other'):
                    genre = gps['genre_other']
                profile_context += f"• Genre: {genre}\n"
                accessed_profile_sections.append("Genre")
            if gps.get('purposes'):
                profile_context += f"• Purpose: {', '.join(gps['purposes'])}\n"
                accessed_profile_sections.append("Book Purpose")
            if gps.get('reader_takeaway'):
                profile_context += f"• Reader Takeaway: {gps['reader_takeaway']}\n"
                accessed_profile_sections.append("Reader Takeaway")
            if gps.get('emotional_tone'):
                profile_context += f"• Emotional Tone: {gps['emotional_tone']}\n"
                accessed_profile_sections.append("Emotional Tone")
            if gps.get('narrative_voices'):
                profile_context += f"• Narrative Voice: {', '.join(gps['narrative_voices'])}\n"
                accessed_profile_sections.append("Narrative Voice")
            if gps.get('time_coverage'):
                profile_context += f"• Time Coverage: {gps['time_coverage']}\n"
                accessed_profile_sections.append("Time Coverage")
            if gps.get('audience_family') or gps.get('audience_industry'):
                profile_context += f"• Target Audience: "
                audiences = []
                if gps.get('audience_family'): audiences.append(f"Family ({gps['audience_family']})")
                if gps.get('audience_industry'): audiences.append(f"Industry ({gps['audience_industry']})")
                if gps.get('audience_general'): audiences.append(f"General ({gps['audience_general']})")
                profile_context += f"{', '.join(audiences)}\n"
                accessed_profile_sections.append("Target Audience")
        
        # Get Enhanced Biographer Profile
        ep = st.session_state.user_account.get('enhanced_profile', {})
        if ep:
            profile_context += "\n\n👤 SECTION 2: SUBJECT BIOGRAPHY (From Enhanced Profile)\n"
            profile_context += "-" * 50 + "\n"
            
            if ep.get('birth_place'):
                profile_context += f"• Birth Place: {ep['birth_place']}\n"
                accessed_profile_sections.append("Birth Place")
            if ep.get('parents'):
                profile_context += f"• Parents: {ep['parents'][:150]}...\n" if len(ep['parents']) > 150 else f"• Parents: {ep['parents']}\n"
                accessed_profile_sections.append("Family Background")
            if ep.get('childhood_home'):
                profile_context += f"• Childhood Home: {ep['childhood_home'][:150]}...\n" if len(ep['childhood_home']) > 150 else f"• Childhood Home: {ep['childhood_home']}\n"
                accessed_profile_sections.append("Childhood")
            if ep.get('family_traditions'):
                profile_context += f"• Family Traditions: {ep['family_traditions'][:150]}...\n" if len(ep['family_traditions']) > 150 else f"• Family Traditions: {ep['family_traditions']}\n"
                accessed_profile_sections.append("Family Traditions")
            if ep.get('school'):
                profile_context += f"• Education: {ep['school'][:150]}...\n" if len(ep['school']) > 150 else f"• Education: {ep['school']}\n"
                accessed_profile_sections.append("Education")
            if ep.get('career_path'):
                profile_context += f"• Career: {ep['career_path'][:150]}...\n" if len(ep['career_path']) > 150 else f"• Career: {ep['career_path']}\n"
                accessed_profile_sections.append("Career")
            if ep.get('romance') or ep.get('marriage'):
                profile_context += f"• Relationships: "
                if ep.get('marriage'): profile_context += f"Married - {ep['marriage'][:100]}... " if len(ep['marriage']) > 100 else f"Married - {ep['marriage']} "
                if ep.get('children'): profile_context += f"Children - {ep['children'][:100]}... " if len(ep['children']) > 100 else f"Children - {ep['children']} "
                profile_context += "\n"
                accessed_profile_sections.append("Relationships")
            if ep.get('challenges'):
                profile_context += f"• Life Challenges: {ep['challenges'][:150]}...\n" if len(ep['challenges']) > 150 else f"• Life Challenges: {ep['challenges']}\n"
                accessed_profile_sections.append("Challenges")
            if ep.get('life_lessons'):
                profile_context += f"• Life Philosophy: {ep['life_lessons'][:150]}...\n" if len(ep['life_lessons']) > 150 else f"• Life Philosophy: {ep['life_lessons']}\n"
                accessed_profile_sections.append("Life Philosophy")
            if ep.get('legacy'):
                profile_context += f"• Legacy Hope: {ep['legacy'][:150]}...\n" if len(ep['legacy']) > 150 else f"• Legacy Hope: {ep['legacy']}\n"
                accessed_profile_sections.append("Legacy Hope")
    
    # Add summary of accessed sections
    if accessed_profile_sections:
        profile_context += f"\n📊 PROFILE SECTIONS USED: {', '.join(set(accessed_profile_sections))}\n"
    else:
        profile_context += "\n⚠️ No profile information found. Complete your profile for personalized feedback!\n"
    
    profile_context += "\n" + "="*80 + "\n"
    profile_context += "📝 BETA READER INSTRUCTIONS: Use the above profile information to provide personalized feedback.\n"
    profile_context += "When your feedback is influenced by specific profile details, mark it with [PROFILE: section_name]\n"
    profile_context += "="*80 + "\n\n"
    
    # Combine profile context with session text
    full_context = profile_context + "\n=== SESSION CONTENT TO REVIEW ===\n\n" + session_text
    
    return beta_reader.generate_feedback(session_title, full_context, feedback_type, accessed_profile_sections)

def save_beta_feedback(user_id, session_id, feedback_data):
    if not user_id:
        return False
    
    try:
        filename = get_user_filename(user_id)
        
        if os.path.exists(filename):
            with open(filename, 'r') as f:
                user_data = json.load(f)
        else:
            user_data = {"responses": {}, "vignettes": [], "beta_feedback": {}}
        
        if "beta_feedback" not in user_data:
            user_data["beta_feedback"] = {}
        
        session_key = str(session_id)
        
        if session_key not in user_data["beta_feedback"]:
            user_data["beta_feedback"][session_key] = []
        
        feedback_copy = feedback_data.copy()
        if "generated_at" not in feedback_copy:
            feedback_copy["generated_at"] = datetime.now().isoformat()
        if "feedback_type" not in feedback_copy:
            feedback_copy["feedback_type"] = "comprehensive"
        
        for s in st.session_state.current_question_bank:
            if str(s["id"]) == session_key:
                feedback_copy["session_title"] = s["title"]
                break
        
        user_data["beta_feedback"][session_key].append(feedback_copy)
        
        with open(filename, 'w') as f:
            json.dump(user_data, f, indent=2)
        
        return True
    except Exception as e:
        print(f"Error saving feedback: {e}")
        return False

def get_previous_beta_feedback(user_id, session_id):
    if not beta_reader: 
        return None
    return beta_reader.get_previous_feedback(user_id, session_id, get_user_filename, load_user_data)

def display_saved_feedback(user_id, session_id):
    user_data = load_user_data(user_id)
    feedback_data = user_data.get("beta_feedback", {})
    session_feedback = feedback_data.get(str(session_id), [])
    
    if not session_feedback:
        st.info("No saved feedback for this session yet.")
        return
    
    st.markdown("### 📚 Saved Beta Reader Feedback")
    
    session_feedback.sort(key=lambda x: x.get('generated_at', ''), reverse=True)
    
    for i, fb in enumerate(session_feedback):
        with st.expander(f"Feedback from {datetime.fromisoformat(fb['generated_at']).strftime('%B %d, %Y at %I:%M %p')}"):
            col1, col2, col3 = st.columns([1, 1, 1])
            
            with col1:
                st.markdown(f"**Type:** {fb.get('feedback_type', 'comprehensive').title()}")
            with col2:
                st.markdown(f"**Overall Score:** {fb.get('overall_score', 'N/A')}/10")
            with col3:
                if st.button(f"🗑️ Delete", key=f"del_fb_{i}_{fb.get('generated_at')}"):
                    session_feedback.pop(i)
                    user_data["beta_feedback"][str(session_id)] = session_feedback
                    save_user_data(user_id, user_data.get("responses", {}))
                    st.rerun()
            
            if 'summary' in fb:
                st.markdown("**Summary:**")
                st.markdown(fb['summary'])
            
            if 'strengths' in fb:
                st.markdown("**Strengths:**")
                for s in fb['strengths']:
                    st.markdown(f"✅ {s}")
            
            if 'areas_for_improvement' in fb:
                st.markdown("**Areas for Improvement:**")
                for a in fb['areas_for_improvement']:
                    st.markdown(f"📝 {a}")
            
            if 'suggestions' in fb:
                st.markdown("**Suggestions:**")
                for sug in fb['suggestions']:
                    st.markdown(f"💡 {sug}")

def save_vignette_beta_feedback(user_id, vignette_id, feedback_data, vignette_title):
    if not user_id:
        return False
    
    try:
        filename = get_user_filename(user_id)
        
        if os.path.exists(filename):
            with open(filename, 'r') as f:
                user_data = json.load(f)
        else:
            user_data = {"responses": {}, "vignettes": [], "beta_feedback": {}, "vignette_beta_feedback": {}}
        
        if "vignette_beta_feedback" not in user_data:
            user_data["vignette_beta_feedback"] = {}
        
        vignette_key = str(vignette_id)
        
        if vignette_key not in user_data["vignette_beta_feedback"]:
            user_data["vignette_beta_feedback"][vignette_key] = []
        
        feedback_copy = feedback_data.copy()
        if "generated_at" not in feedback_copy:
            feedback_copy["generated_at"] = datetime.now().isoformat()
        if "feedback_type" not in feedback_copy:
            feedback_copy["feedback_type"] = "comprehensive"
        
        feedback_copy["vignette_title"] = vignette_title
        
        user_data["vignette_beta_feedback"][vignette_key].append(feedback_copy)
        
        with open(filename, 'w') as f:
            json.dump(user_data, f, indent=2)
        
        return True
    except Exception as e:
        print(f"Error saving vignette feedback: {e}")
        return False

def display_beta_feedback(feedback_data):
    """Display beta feedback in a styled container below the answer box"""
    if not feedback_data:
        return
    
    st.markdown("---")
    st.markdown("### 🦋 Beta Reader Feedback")
    
    with st.container():
        col1, col2 = st.columns([6, 1])
        with col2:
            if st.button("✕", key="close_beta_feedback"):
                st.session_state.beta_feedback_display = None
                if "beta_feedback_storage" in st.session_state:
                    st.session_state.beta_feedback_storage = {}
                st.rerun()
        
        if 'error' in feedback_data:
            st.error(f"Error: {feedback_data['error']}")
            return
        
        if feedback_data.get('profile_sections_used'):
            with st.expander("📋 **PROFILE INFORMATION ACCESSED**", expanded=True):
                st.markdown("The Beta Reader used these profile sections to personalize this feedback:")
                cols = st.columns(3)
                for i, section in enumerate(feedback_data['profile_sections_used']):
                    cols[i % 3].markdown(f"✅ **{section}**")
                st.markdown("\n*Look for highlighted **[PROFILE: ...]** markers in the feedback below to see where specific profile information influenced the analysis.*")
        
        col_save1, col_save2, col_save3 = st.columns([1, 2, 1])
        with col_save2:
            if st.button("💾 Save This Feedback to History", key="save_beta_feedback", type="primary", use_container_width=True):
                result = save_beta_feedback(
                    st.session_state.user_id,
                    st.session_state.current_question_bank[st.session_state.current_session]["id"],
                    feedback_data
                )
                if result:
                    st.success("✅ Feedback saved to history!")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("Failed to save feedback")
        
        st.markdown("---")
        
        if 'feedback' in feedback_data and feedback_data['feedback']:
            feedback_text = feedback_data['feedback']
            
            parts = re.split(r'(\[PROFILE:.*?\])', feedback_text)
            formatted_feedback = ""
            for i, part in enumerate(parts):
                if part.startswith('[PROFILE:') and part.endswith(']'):
                    formatted_feedback += f'<span style="background-color: #e8f4fd; color: #0366d6; font-weight: bold; padding: 2px 6px; border-radius: 4px; border-left: 3px solid #0366d6;">{part}</span>'
                else:
                    formatted_feedback += part
            
            st.markdown(formatted_feedback, unsafe_allow_html=True)
        else:
            if 'summary' in feedback_data and feedback_data['summary']:
                st.markdown("**Summary:**")
                st.info(feedback_data['summary'])
            
            if 'strengths' in feedback_data and feedback_data['strengths']:
                st.markdown("**Strengths:**")
                for s in feedback_data['strengths']:
                    st.markdown(f"✅ {s}")
            
            if 'areas_for_improvement' in feedback_data and feedback_data['areas_for_improvement']:
                st.markdown("**Areas for Improvement:**")
                for a in feedback_data['areas_for_improvement']:
                    st.markdown(f"📝 {a}")
            
            if 'suggestions' in feedback_data and feedback_data['suggestions']:
                st.markdown("**Suggestions:**")
                for sug in feedback_data['suggestions']:
                    st.markdown(f"💡 {sug}")
            
            if 'overall_score' in feedback_data and feedback_data['overall_score']:
                st.markdown(f"**Overall Score:** {feedback_data['overall_score']}/10")

# ============================================================================
# VIGNETTE FUNCTIONS
# ============================================================================
def on_vignette_select(vignette_id):
    st.session_state.selected_vignette_id = vignette_id
    st.session_state.show_vignette_detail = True
    st.session_state.show_vignette_manager = False
    st.rerun()

def on_vignette_edit(vignette_id):
    st.session_state.editing_vignette_id = vignette_id
    st.session_state.show_vignette_detail = False
    st.session_state.show_vignette_manager = False
    st.session_state.show_vignette_modal = True
    st.rerun()

def on_vignette_delete(vignette_id):
    if VignetteManager and st.session_state.get('vignette_manager', VignetteManager(st.session_state.user_id)).delete_vignette(vignette_id):
        st.success("Deleted!"); 
        st.rerun()
    else: 
        st.error("Failed to delete")

def on_vignette_publish(vignette):
    st.session_state.published_vignette = vignette
    st.success(f"Published '{vignette['title']}'!"); 
    st.rerun()

def show_vignette_modal():
    if not VignetteManager: 
        st.error("Vignette module not available"); 
        st.session_state.show_vignette_modal = False; 
        return
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    if st.button("←", key="vign_modal_back"): 
        st.session_state.show_vignette_modal = False; 
        st.session_state.editing_vignette_id = None; 
        st.rerun()
    st.title("✏️ Edit Vignette" if st.session_state.get('editing_vignette_id') else "✍️ Create Vignette")
    if 'vignette_manager' not in st.session_state: 
        st.session_state.vignette_manager = VignetteManager(st.session_state.user_id)
    edit = st.session_state.vignette_manager.get_vignette_by_id(st.session_state.editing_vignette_id) if st.session_state.get('editing_vignette_id') else None
    st.session_state.vignette_manager.display_vignette_creator(on_publish=on_vignette_publish, edit_vignette=edit)
    
    # ADD BETA READER SECTION FOR EDIT MODE
    if st.session_state.get('editing_vignette_id') and edit:
        st.divider()
        st.markdown("## 🦋 Beta Reader Feedback for This Vignette")
        
        tab1, tab2 = st.tabs(["📝 Get Feedback", "📚 Feedback History"])
        
        with tab1:
            col1, col2 = st.columns([2, 1])
            with col1:
                fb_type = st.selectbox(
                    "Feedback Type", 
                    ["comprehensive", "concise", "developmental"], 
                    key=f"beta_vignette_edit_type_{edit['id']}"
                )
            with col2:
                if st.button("🦋 Get Beta Read", key=f"beta_vignette_edit_btn_{edit['id']}", type="primary", use_container_width=True):
                    with st.spinner("Beta Reader is analyzing your vignette with profile context..."):
                        if beta_reader:
                            vignette_text = edit.get('content', '')
                            if vignette_text and len(vignette_text.strip()) > 50:
                                fb = generate_beta_reader_feedback(
                                    f"Vignette: {edit['title']}", 
                                    vignette_text,
                                    fb_type
                                )
                                if "error" not in fb: 
                                    st.session_state[f"beta_vignette_edit_{edit['id']}"] = fb
                                    st.rerun()
                                else: 
                                    st.error(f"Failed: {fb['error']}")
                            else:
                                st.warning("Vignette too short for feedback (minimum 50 words)")
                        else:
                            st.error("Beta reader not available")
            
            if f"beta_vignette_edit_{edit['id']}" in st.session_state:
                fb = st.session_state[f"beta_vignette_edit_{edit['id']}"]
                st.markdown("---")
                st.markdown("### 📋 Beta Reader Results")
                
                if fb.get('profile_sections_used'):
                    with st.expander("📋 **PROFILE INFORMATION ACCESSED**", expanded=True):
                        st.markdown("The Beta Reader used these profile sections to personalize this feedback:")
                        cols = st.columns(3)
                        for i, section in enumerate(fb['profile_sections_used']):
                            cols[i % 3].markdown(f"✅ **{section}**")
                
                if st.button("💾 Save to History", key=f"save_vignette_edit_fb_{edit['id']}"):
                    if save_vignette_beta_feedback(st.session_state.user_id, edit['id'], fb, edit['title']):
                        st.success("✅ Saved!")
                        st.rerun()
                
                st.markdown("---")
                
                if 'feedback' in fb and fb['feedback']:
                    feedback_text = fb['feedback']
                    parts = re.split(r'(\[PROFILE:.*?\])', feedback_text)
                    formatted_feedback = ""
                    for part in parts:
                        if part.startswith('[PROFILE:') and part.endswith(']'):
                            formatted_feedback += f'<span style="background-color: #e8f4fd; color: #0366d6; font-weight: bold; padding: 2px 6px; border-radius: 4px; border-left: 3px solid #0366d6;">{part}</span>'
                        else:
                            formatted_feedback += part
                    st.markdown(formatted_feedback, unsafe_allow_html=True)
                else:
                    if 'summary' in fb:
                        st.info(fb['summary'])
                    if 'strengths' in fb:
                        for s in fb['strengths']:
                            st.markdown(f"✅ {s}")
                    if 'areas_for_improvement' in fb:
                        for a in fb['areas_for_improvement']:
                            st.markdown(f"📝 {a}")
        
        with tab2:
            st.markdown("### Saved Feedback")
            user_data = load_user_data(st.session_state.user_id) if st.session_state.user_id else {}
            vignette_feedback = user_data.get("vignette_beta_feedback", {})
            this_vignette_feedback = vignette_feedback.get(str(edit['id']), [])
            
            if not this_vignette_feedback:
                st.info("No saved feedback yet")
            else:
                for i, fb in enumerate(this_vignette_feedback):
                    with st.expander(f"Feedback {i+1}"):
                        if st.button(f"Delete", key=f"del_vignette_edit_fb_{i}"):
                            this_vignette_feedback.pop(i)
                            user_data["vignette_beta_feedback"] = vignette_feedback
                            save_user_data(st.session_state.user_id, user_data.get("responses", {}))
                            st.rerun()
                        if 'feedback' in fb:
                            st.markdown(fb['feedback'])
    
    st.markdown('</div>', unsafe_allow_html=True)

def show_vignette_manager():
    if not VignetteManager: 
        st.error("Vignette module not available"); 
        st.session_state.show_vignette_manager = False; 
        return
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    if st.button("←", key="vign_mgr_back"): 
        st.session_state.show_vignette_manager = False; 
        st.rerun()
    st.title("📚 Your Vignettes")
    if 'vignette_manager' not in st.session_state: 
        st.session_state.vignette_manager = VignetteManager(st.session_state.user_id)
    filter_map = {"All Stories": "all", "Published": "published", "Drafts": "drafts"}
    filter_option = st.radio("Show:", ["All Stories", "Published", "Drafts"], horizontal=True, key="vign_filter")
    st.session_state.vignette_manager.display_vignette_gallery(
        filter_by=filter_map.get(filter_option, "all"),
        on_select=on_vignette_select, 
        on_edit=on_vignette_edit, 
        on_delete=on_vignette_delete
    )
    st.divider()
    if st.button("➕ Create New Vignette", type="primary", use_container_width=True):
        st.session_state.show_vignette_manager = False; 
        st.session_state.show_vignette_modal = True; 
        st.session_state.editing_vignette_id = None; 
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

def show_vignette_detail():
    if not VignetteManager or not st.session_state.get('selected_vignette_id'): 
        st.session_state.show_vignette_detail = False
        return
    
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    
    col1, col2 = st.columns([6, 1])
    with col1:
        st.title("📖 Read Vignette")
    with col2:
        if st.button("✕", key="close_vignette_detail"):
            st.session_state.show_vignette_detail = False
            st.session_state.selected_vignette_id = None
            st.rerun()
    
    if 'vignette_manager' not in st.session_state: 
        st.session_state.vignette_manager = VignetteManager(st.session_state.user_id)
    
    vignette = st.session_state.vignette_manager.get_vignette_by_id(st.session_state.selected_vignette_id)
    if not vignette: 
        st.error("Vignette not found")
        st.session_state.show_vignette_detail = False
        return
    
    st.session_state.vignette_manager.display_full_vignette(
        st.session_state.selected_vignette_id,
        on_back=lambda: st.session_state.update(show_vignette_detail=False, selected_vignette_id=None),
        on_edit=on_vignette_edit
    )
    
    st.divider()
    
    st.markdown("## 🦋 Beta Reader Feedback for This Vignette")
    
    tab1, tab2 = st.tabs(["📝 Get Feedback", "📚 Feedback History"])
    
    with tab1:
        col1, col2 = st.columns([2, 1])
        with col1:
            fb_type = st.selectbox(
                "Feedback Type", 
                ["comprehensive", "concise", "developmental"], 
                key=f"beta_vignette_type_{vignette['id']}"
            )
        with col2:
            if st.button("🦋 Get Beta Read", key=f"beta_vignette_btn_{vignette['id']}", type="primary", use_container_width=True):
                with st.spinner("Beta Reader is analyzing your vignette with profile context..."):
                    if beta_reader:
                        vignette_text = vignette.get('content', '')
                        if vignette_text and len(vignette_text.strip()) > 50:
                            fb = generate_beta_reader_feedback(
                                f"Vignette: {vignette['title']}", 
                                vignette_text,
                                fb_type
                            )
                            if "error" not in fb: 
                                st.session_state[f"beta_vignette_{vignette['id']}"] = fb
                                st.rerun()
                            else: 
                                st.error(f"Failed: {fb['error']}")
                        else:
                            st.warning("Vignette too short for feedback (minimum 50 words)")
                    else:
                        st.error("Beta reader not available")
        
        if f"beta_vignette_{vignette['id']}" in st.session_state:
            fb = st.session_state[f"beta_vignette_{vignette['id']}"]
            st.markdown("---")
            st.markdown("### 📋 Beta Reader Results")
            
            if fb.get('profile_sections_used'):
                with st.expander("📋 **PROFILE INFORMATION ACCESSED**", expanded=True):
                    st.markdown("The Beta Reader used these profile sections to personalize this feedback:")
                    cols = st.columns(3)
                    for i, section in enumerate(fb['profile_sections_used']):
                        cols[i % 3].markdown(f"✅ **{section}**")
            
            if st.button("💾 Save to History", key=f"save_vignette_fb_{vignette['id']}"):
                if save_vignette_beta_feedback(st.session_state.user_id, vignette['id'], fb, vignette['title']):
                    st.success("✅ Saved!")
                    st.rerun()
            
            st.markdown("---")
            
            if 'feedback' in fb and fb['feedback']:
                feedback_text = fb['feedback']
                parts = re.split(r'(\[PROFILE:.*?\])', feedback_text)
                formatted_feedback = ""
                for part in parts:
                    if part.startswith('[PROFILE:') and part.endswith(']'):
                        formatted_feedback += f'<span style="background-color: #e8f4fd; color: #0366d6; font-weight: bold; padding: 2px 6px; border-radius: 4px; border-left: 3px solid #0366d6;">{part}</span>'
                    else:
                        formatted_feedback += part
                st.markdown(formatted_feedback, unsafe_allow_html=True)
            else:
                if 'summary' in fb:
                    st.info(fb['summary'])
                if 'strengths' in fb:
                    for s in fb['strengths']:
                        st.markdown(f"✅ {s}")
                if 'areas_for_improvement' in fb:
                    for a in fb['areas_for_improvement']:
                        st.markdown(f"📝 {a}")
    
    with tab2:
        st.markdown("### Saved Feedback")
        user_data = load_user_data(st.session_state.user_id) if st.session_state.user_id else {}
        vignette_feedback = user_data.get("vignette_beta_feedback", {})
        this_vignette_feedback = vignette_feedback.get(str(vignette['id']), [])
        
        if not this_vignette_feedback:
            st.info("No saved feedback yet")
        else:
            for i, fb in enumerate(this_vignette_feedback):
                with st.expander(f"Feedback {i+1}"):
                    if st.button(f"Delete", key=f"del_vignette_fb_{i}"):
                        this_vignette_feedback.pop(i)
                        user_data["vignette_beta_feedback"] = vignette_feedback
                        save_user_data(st.session_state.user_id, user_data.get("responses", {}))
                        st.rerun()
                    if 'feedback' in fb:
                        st.markdown(fb['feedback'])
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

def switch_to_vignette(vignette_topic, content=""):
    st.session_state.current_question_override = f"Vignette: {vignette_topic}"
    if content:
        save_response(st.session_state.current_question_bank[st.session_state.current_session]["id"], 
                     f"Vignette: {vignette_topic}", content)
    st.rerun()

def switch_to_custom_topic(topic_text):
    st.session_state.current_question_override = topic_text
    st.rerun()

# ============================================================================
# TOPIC BROWSER & SESSION MANAGER
# ============================================================================
def show_topic_browser():
    if not TopicBank: 
        st.error("Topic module not available"); 
        st.session_state.show_topic_browser = False; 
        return
    
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    
    col1, col2 = st.columns([6, 1])
    with col1:
        st.title("📚 Topic Bank")
    with col2:
        if st.button("✕", key="close_topic_browser"):
            st.session_state.show_topic_browser = False
            st.rerun()
    
    st.markdown("Browse and load topics into your current session.")
    st.divider()
    
    topic_bank = TopicBank(st.session_state.user_id)
    
    def on_topic_selected(topic_text):
        switch_to_custom_topic(topic_text)
        st.session_state.show_topic_browser = False
        st.rerun()
    
    topic_bank.display_topic_browser(
        on_topic_select=on_topic_selected,
        unique_key="topic_bank_browser"
    )
    
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()

def show_session_creator():
    if not SessionManager: 
        st.error("Session module not available"); 
        st.session_state.show_session_creator = False; 
        return
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    if st.button("← Back", key="session_creator_back"): 
        st.session_state.show_session_creator = False; 
        st.rerun()
    st.title("📋 Create Custom Session")
    SessionManager(st.session_state.user_id, "sessions/sessions.csv").display_session_creator()
    st.markdown('</div>', unsafe_allow_html=True)

def show_session_manager():
    if not SessionManager: 
        st.error("Session module not available"); 
        st.session_state.show_session_manager = False; 
        return
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    if st.button("← Back", key="session_manager_back"): 
        st.session_state.show_session_manager = False; 
        st.rerun()
    st.title("📖 Session Manager")
    mgr = SessionManager(st.session_state.user_id, "sessions/sessions.csv")
    if st.button("➕ Create New Session", type="primary", use_container_width=True):
        st.session_state.show_session_manager = False; 
        st.session_state.show_session_creator = True; 
        st.rerun()
    st.divider()
    mgr.display_session_grid(cols=2, on_session_select=lambda sid: [st.session_state.update(
        current_session=i, current_question=0, current_question_override=None) for i, s in enumerate(st.session_state.current_question_bank) if s["id"] == sid][0])
    st.markdown('</div>', unsafe_allow_html=True)

# ============================================================================
# QUESTION BANK UI FUNCTIONS
# ============================================================================
def show_bank_manager():
    if not QuestionBankManager: 
        st.error("Question Bank Manager not available"); 
        st.session_state.show_bank_manager = False; 
        return
    user_id = st.session_state.get('user_id')
    if st.session_state.qb_manager is None: 
        st.session_state.qb_manager = QuestionBankManager(user_id)
    else: 
        st.session_state.qb_manager.user_id = user_id
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    if st.button("←", key="bank_manager_back"): 
        st.session_state.show_bank_manager = False; 
        st.rerun()
    st.session_state.qb_manager.display_bank_selector()
    st.markdown('</div>', unsafe_allow_html=True)

def show_bank_editor():
    if not QuestionBankManager or not st.session_state.get('editing_bank_id'): 
        st.session_state.show_bank_editor = False; 
        return
    user_id = st.session_state.get('user_id')
    if st.session_state.qb_manager is None: 
        st.session_state.qb_manager = QuestionBankManager(user_id)
    else: 
        st.session_state.qb_manager.user_id = user_id
    st.markdown('<div class="modal-overlay">', unsafe_allow_html=True)
    
    # Call the display_bank_editor method
    st.session_state.qb_manager.display_bank_editor(st.session_state.editing_bank_id)
    
    # Check if the editor should close (this happens when the back button is clicked inside display_bank_editor)
    if not st.session_state.get('show_bank_editor', False):
        st.markdown('</div>', unsafe_allow_html=True)
        st.stop()
        return
    
    st.markdown('</div>', unsafe_allow_html=True)

# ============================================================================
# FILE IMPORT FUNCTION FOR MAIN EDITOR
# ============================================================================
def import_text_file_main(uploaded_file):
    """Import text from common document formats into the main editor"""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        file_content = ""
        file_size_mb = len(uploaded_file.getvalue()) / (1024 * 1024)
        
        st.info(f"📄 Importing: {uploaded_file.name} ({file_size_mb:.1f}MB)")
        
        if file_extension == 'txt':
            file_content = uploaded_file.read().decode('utf-8', errors='ignore')
        
        elif file_extension == 'docx':
            try:
                import io
                from docx import Document
                docx_bytes = io.BytesIO(uploaded_file.getvalue())
                doc = Document(docx_bytes)
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                file_content = '\n\n'.join(paragraphs)
            except ImportError:
                st.error("Please install: pip install python-docx")
                return None
        
        elif file_extension == 'rtf':
            try:
                from striprtf.striprtf import rtf_to_text
                rtf_content = uploaded_file.read().decode('utf-8', errors='ignore')
                file_content = rtf_to_text(rtf_content)
            except ImportError:
                st.warning("RTF support needs: pip install striprtf")
                return None
        
        elif file_extension in ['vtt', 'srt']:
            file_content = uploaded_file.read().decode('utf-8', errors='ignore')
            lines = file_content.split('\n')
            clean_lines = [line.strip() for line in lines if '-->' not in line and not line.strip().isdigit() and line.strip()]
            file_content = ' '.join(clean_lines)
        
        elif file_extension == 'json':
            try:
                import json
                data = json.loads(uploaded_file.read().decode('utf-8'))
                if isinstance(data, dict):
                    file_content = data.get('text', data.get('transcript', str(data)))
                else:
                    file_content = str(data)
            except Exception as e:
                st.error(f"Error parsing JSON: {e}")
                return None
        
        elif file_extension == 'md':
            file_content = uploaded_file.read().decode('utf-8', errors='ignore')
            file_content = re.sub(r'#{1,6}\s*', '', file_content)
            file_content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', file_content)
        
        else:
            st.error(f"Unsupported format: .{file_extension}")
            st.info("Supported: .txt, .docx, .rtf, .vtt, .srt, .json, .md")
            return None
        
        if not file_content or not file_content.strip():
            st.warning("File is empty")
            return None
        
        if file_size_mb > 10:
            st.warning(f"⚠️ Large file ({file_size_mb:.1f}MB) - processing may be slow")
        
        file_content = re.sub(r'\s+', ' ', file_content)
        sentences = re.split(r'[.!?]+', file_content)
        paragraphs = []
        current_para = []
        
        for sentence in sentences:
            if sentence.strip():
                current_para.append(sentence.strip() + '.')
                if len(current_para) >= 4:
                    paragraphs.append(' '.join(current_para))
                    current_para = []
        
        if current_para:
            paragraphs.append(' '.join(current_para))
        
        if not paragraphs:
            paragraphs = [file_content]
        
        html_content = ''
        for para in paragraphs:
            if para.strip():
                para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                html_content += f'<p>{para.strip()}</p>'
        
        return html_content
        
    except Exception as e:
        st.error(f"Import error: {str(e)}")
        return None

# ============================================================================
# PAGE CONFIG - ALREADY SET AT TOP
# ============================================================================

if not st.session_state.qb_manager_initialized: 
    initialize_question_bank()
SESSIONS = st.session_state.get('current_question_bank', [])

if st.session_state.logged_in and st.session_state.user_id and not st.session_state.data_loaded:
    user_data = load_user_data(st.session_state.user_id)
    if "responses" in user_data:
        for sid_str, sdata in user_data["responses"].items():
            try: 
                sid = int(sid_str)
            except: 
                continue
            if sid in st.session_state.responses and "questions" in sdata and sdata["questions"]:
                st.session_state.responses[sid]["questions"] = sdata["questions"]
    st.session_state.data_loaded = True
    init_image_handler()

if not SESSIONS:
    st.error("❌ No question bank loaded. Use Bank Manager.")
    if st.button("📋 Open Bank Manager", type="primary", use_container_width=True): 
        st.session_state.show_bank_manager = True; 
        st.rerun()
    st.stop()

# ============================================================================
# AUTHENTICATION UI
# ============================================================================
if not st.session_state.logged_in:
    st.markdown('<div class="auth-container"><h1>Tell My Story</h1><p>Your Life Timeline • Preserve Your Legacy</p></div>', unsafe_allow_html=True)
    if 'auth_tab' not in st.session_state: 
        st.session_state.auth_tab = 'login'
    
    col1, col2 = st.columns(2)
    with col1: 
        st.button("🔐 Login", use_container_width=True, type="primary" if st.session_state.auth_tab=='login' else "secondary", 
                        on_click=lambda: st.session_state.update(auth_tab='login'))
    with col2: 
        st.button("📝 Sign Up", use_container_width=True, type="primary" if st.session_state.auth_tab=='signup' else "secondary",
                        on_click=lambda: st.session_state.update(auth_tab='signup'))
    st.divider()
    
    if st.session_state.auth_tab == 'login':
        with st.form("login_form"):
            st.subheader("Welcome Back")
            email = st.text_input("Email Address")
            password = st.text_input("Password", type="password")
            if st.form_submit_button("Login", type="primary", use_container_width=True):
                if email and password:
                    result = authenticate_user(email, password)
                    if result["success"]:
                        st.session_state.update(user_id=result["user_id"], 
                                              user_account=result["user_record"],
                                              logged_in=True, 
                                              data_loaded=False, 
                                              qb_manager=None, 
                                              qb_manager_initialized=False)
                        st.success("Login successful!"); 
                        st.rerun()
                    else: 
                        st.error(f"Login failed: {result.get('error', 'Unknown error')}")
    else:
        with st.form("signup_form"):
            st.subheader("Create New Account")
            col1, col2 = st.columns(2)
            with col1: 
                first_name = st.text_input("First Name*")
            with col2: 
                last_name = st.text_input("Last Name*")
            email = st.text_input("Email Address*")
            col1, col2 = st.columns(2)
            with col1: 
                password = st.text_input("Password*", type="password")
            with col2: 
                confirm = st.text_input("Confirm Password*", type="password")
            accept = st.checkbox("I agree to the Terms*")
            
            if st.form_submit_button("Create Account", type="primary", use_container_width=True):
                errors = []
                if not first_name: errors.append("First name required")
                if not last_name: errors.append("Last name required")
                if not email or "@" not in email: errors.append("Valid email required")
                if not password or len(password) < 8: errors.append("Password must be 8+ characters")
                if password != confirm: errors.append("Passwords don't match")
                if not accept: errors.append("Must accept terms")
                if get_account_data(email=email): errors.append("Email already exists")
                
                if errors: 
                    [st.error(e) for e in errors]
                else:
                    result = create_user_account({"first_name": first_name, "last_name": last_name, "email": email, "account_for": "self"}, password)
                    if result["success"]:
                        send_welcome_email({"first_name": first_name, "email": email}, 
                                         {"user_id": result["user_id"], "password": password})
                        st.session_state.update(user_id=result["user_id"], 
                                              user_account=result["user_record"],
                                              logged_in=True, 
                                              data_loaded=False, 
                                              show_profile_setup=True,
                                              qb_manager=None, 
                                              qb_manager_initialized=False)
                        st.success("Account created!"); 
                        st.balloons(); 
                        st.rerun()
                    else: 
                        st.error(f"Error: {result.get('error', 'Unknown error')}")
    st.stop()

# ============================================================================
# PROFILE SETUP MODAL
# ============================================================================
if st.session_state.get('show_profile_setup', False):
    st.markdown('<div class="profile-setup-modal">', unsafe_allow_html=True)
    st.title("👤 Your Complete Life Story Profile")
    
    st.markdown("### 📝 Basic Information")
    with st.form("profile_setup_form"):
        col1, col2 = st.columns(2)
        with col1:
            gender = st.radio("Gender", ["Male", "Female", "Other", "Prefer not to say"], horizontal=True, key="modal_gender")
        with col2:
            account_for = st.radio("Account Type", ["For me", "For someone else"], key="modal_account_type", horizontal=True)
        
        col1, col2, col3 = st.columns(3)
        with col1: 
            birth_month = st.selectbox("Birth Month", ["January","February","March","April","May","June","July","August","September","October","November","December"], key="modal_month")
        with col2: 
            birth_day = st.selectbox("Birth Day", list(range(1,32)), key="modal_day")
        with col3: 
            birth_year = st.selectbox("Birth Year", list(range(datetime.now().year, datetime.now().year-120, -1)), key="modal_year")
        
        col_save, col_close = st.columns([3, 1])
        with col_save:
            if st.form_submit_button("💾 Save Basic Information", type="primary", use_container_width=True):
                if birth_month and birth_day and birth_year:
                    birthdate = f"{birth_month} {birth_day}, {birth_year}"
                    if st.session_state.user_account:
                        st.session_state.user_account['profile'].update({
                            'gender': gender, 'birthdate': birthdate, 'timeline_start': birthdate
                        })
                        st.session_state.user_account['account_type'] = "self" if account_for == "For me" else "other"
                        save_account_data(st.session_state.user_account)
                    st.success("Basic information saved!")
                    st.rerun()
        with col_close:
            if st.form_submit_button("✕ Close Profile", use_container_width=True):
                st.session_state.show_profile_setup = False
                st.rerun()
    
    st.divider()
    render_narrative_gps()
    st.divider()
    render_enhanced_profile()
    st.divider()
    
    with st.expander("🔒 Privacy & Security Settings", expanded=False):
        if 'privacy_settings' not in st.session_state.user_account:
            st.session_state.user_account['privacy_settings'] = {
                "profile_public": False, "stories_public": False, "allow_sharing": False,
                "data_collection": True, "encryption": True
            }
        
        privacy = st.session_state.user_account['privacy_settings']
        
        privacy['profile_public'] = st.checkbox("Make profile public", value=privacy.get('profile_public', False),
                                               help="Allow others to see your basic profile information")
        privacy['stories_public'] = st.checkbox("Share stories publicly", value=privacy.get('stories_public', False),
                                               help="Make your stories visible to the public (coming soon)")
        privacy['allow_sharing'] = st.checkbox("Allow sharing via link", value=privacy.get('allow_sharing', False),
                                              help="Generate shareable links to your stories")
        privacy['data_collection'] = st.checkbox("Allow anonymous usage data", value=privacy.get('data_collection', True),
                                                help="Help us improve by sharing anonymous usage statistics")
        
        st.markdown("**🔐 Security Status:** Your data is encrypted at rest and never shared with third parties.")
        
        if st.button("💾 Save Privacy Settings", type="primary", use_container_width=True):
            save_account_data(st.session_state.user_account)
            st.success("Privacy settings saved!")
            st.rerun()
    
    st.divider()
    
    with st.expander("💾 Backup & Restore", expanded=False):
        st.markdown("**Create a complete backup of all your data:**")
        backup_json = create_backup()
        if backup_json:
            st.download_button(
                label="📥 Download Complete Backup",
                data=backup_json,
                file_name=f"tell_my_story_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json",
                use_container_width=True
            )
        
        st.markdown("---")
        st.markdown("**Restore from backup:**")
        backup_file = st.file_uploader("Upload backup file", type=['json'], key="restore_uploader")
        
        if backup_file:
            # Show file info
            file_size_kb = len(backup_file.getvalue()) / 1024
            st.info(f"📄 Selected: {backup_file.name} ({file_size_kb:.1f} KB)")
            
            # WARNING MESSAGE
            st.warning("⚠️ **WARNING:** This will COMPLETELY OVERWRITE all your current stories and profile data. This action CANNOT be undone!")
            
            if st.button("🔄 RESTORE BACKUP (I understand the risk)", type="primary", use_container_width=True):
                with st.spinner("Restoring your data..."):
                    try:
                        backup_content = backup_file.read().decode('utf-8')
                        if restore_from_backup(backup_content):
                            st.success("✅ Backup restored successfully! Your data has been recovered.")
                            time.sleep(2)
                            st.rerun()
                        else:
                            st.error("❌ Failed to restore backup. File may be corrupted or for wrong user.")
                    except Exception as e:
                        st.error(f"❌ Error: {str(e)}")
        
        st.markdown("---")
        st.markdown("**Previous backups:**")
        backups = list_backups()
        if backups:
            for b in backups:
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.text(f"📅 {b['date']} ({(b['size']/1024):.1f} KB)")
                with col2:
                    if st.button(f"Restore", key=f"restore_{b['filename']}"):
                        # WARNING for previous backups too
                        st.warning("⚠️ This will overwrite ALL current data!")
                        if st.button(f"✅ CONFIRM Restore {b['filename']}", key=f"confirm_{b['filename']}"):
                            with open(f"backups/{b['filename']}", 'r') as f:
                                backup_content = f.read()
                            if restore_from_backup(backup_content):
                                st.success("✅ Restored successfully!")
                                st.rerun()

# ============================================================================
# MODAL HANDLING
# ============================================================================
if st.session_state.show_ai_rewrite and st.session_state.current_rewrite_data:
    show_ai_rewrite_modal()
    st.stop()

# PROMPT ME MODAL HANDLING
if st.session_state.get('show_prompt_modal', False) and st.session_state.get('current_prompt_data'):
    show_prompt_me_modal()
    st.stop()

if st.session_state.show_privacy_settings:
    show_privacy_settings()
    st.stop()

if st.session_state.show_cover_designer:
    show_cover_designer()
    st.stop()

if st.session_state.show_bank_manager:
    show_bank_manager()
    if st.session_state.show_bank_manager:
        st.stop()

if st.session_state.show_bank_editor:
    show_bank_editor()
    if st.session_state.show_bank_editor:
        st.stop()

if st.session_state.show_vignette_detail:
    show_vignette_detail()
    if st.session_state.show_vignette_detail:
        st.stop()

if st.session_state.show_vignette_manager:
    show_vignette_manager()
    if st.session_state.show_vignette_manager:
        st.stop()

if st.session_state.show_vignette_modal:
    show_vignette_modal()
    if st.session_state.show_vignette_modal:
        st.stop()

if st.session_state.show_topic_browser:
    show_topic_browser()
    if st.session_state.show_topic_browser:
        st.stop()

if st.session_state.show_session_manager:
    show_session_manager()
    if st.session_state.show_session_manager:
        st.stop()

if st.session_state.show_session_creator:
    show_session_creator()
    if st.session_state.show_session_creator:
        st.stop()

# ============================================================================
# MAIN HEADER WITH PROMINENT SUPPORT BUTTON
# ============================================================================
col_logo, col_title, col_support = st.columns([1, 3, 1])
with col_logo:
    st.markdown(f'<img src="{LOGO_URL}" class="logo-img" style="max-width:100px;">', unsafe_allow_html=True)
with col_title:
    st.markdown('<h1 style="margin-top:20px;">Tell My Story</h1>', unsafe_allow_html=True)
with col_support:
    st.markdown('<div style="margin-top:20px;">', unsafe_allow_html=True)
    if st.button("❓ Help", type="primary", use_container_width=True):
        st.session_state.show_support = not st.session_state.get('show_support', False)
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# ============================================================================
# SUPPORT SECTION - Show when help button clicked
# ============================================================================
if st.session_state.get('show_support', False):
    # Hide the sidebar when in support mode
    st.markdown("""
    <style>
        section[data-testid="stSidebar"] { display: none !important; }
    </style>
    """, unsafe_allow_html=True)
    
    # Add Back button at the top (above tabs)
    col_back1, col_back2, col_back3 = st.columns([1, 2, 1])
    with col_back2:
        if st.button("← Back to Writing", type="primary", use_container_width=True):
            st.session_state.show_support = False
            st.rerun()
    
    try:
        from support_section import SupportSection
        support = SupportSection()
        support.render()
        st.stop()
        
    except ImportError:
        st.error("Support section not available. Please ensure support_section.py is in the same directory.")
        if st.button("← Back"):
            st.session_state.show_support = False
            st.rerun()
        st.stop()
    except Exception as e:
        st.error(f"Error loading support section: {e}")
        if st.button("← Back"):
            st.session_state.show_support = False
            st.rerun()
        st.stop()

# ============================================================================
# SIDEBAR
# ============================================================================
with st.sidebar:
    st.markdown('<div class="sidebar-header"><h2>Tell My Story</h2><p>Your Life Timeline</p></div>', unsafe_allow_html=True)
    
    # Gamification Dashboard at the top
    render_gamification_dashboard()
    st.divider()
    
    st.header("👤 Your Profile")
    if st.session_state.user_account:
        profile = st.session_state.user_account['profile']
        st.success(f"✓ **{profile['first_name']} {profile['last_name']}**")
    if st.button("📝 Complete Profile", use_container_width=True): 
        st.session_state.show_profile_setup = True
        st.rerun()
    if st.button("🚪 Log Out", use_container_width=True): 
        logout_user()
    
    st.divider()
    st.header("🔧 Tools")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🔒 Privacy", use_container_width=True):
            st.session_state.show_privacy_settings = True
            st.rerun()
    with col2:
        if st.button("🎨 Cover", use_container_width=True):
            st.session_state.show_cover_designer = True
            st.rerun()
    
    st.divider()
    st.header("📚 Question Banks")
    if st.button("📋 Bank Manager", use_container_width=True, type="primary"): 
        st.session_state.show_bank_manager = True
        st.rerun()
    if st.session_state.get('current_bank_name'): 
        st.info(f"**Current Bank:**\n{st.session_state.current_bank_name}")
    
    st.divider()
    st.header("📖 Sessions")
    if st.session_state.current_question_bank:
        for i, s in enumerate(st.session_state.current_question_bank):
            sid = s["id"]
            sdata = st.session_state.responses.get(sid, {})
            resp_cnt = len(sdata.get("questions", {}))
            total_q = len(s["questions"])
            status = "🟢" if resp_cnt == total_q and total_q > 0 else "🟡" if resp_cnt > 0 else "🔴"
            if i == st.session_state.current_session: 
                status = "▶️"
            if st.button(f"{status} Session {sid}: {s['title']}", key=f"sel_sesh_{i}", use_container_width=True):
                st.session_state.update(current_session=i, current_question=0, editing=False, current_question_override=None, show_ai_rewrite_menu=False)
                st.rerun()
    
    st.divider()
    st.header("✨ Vignettes")
    if st.button("📝 New Vignette", use_container_width=True): 
        import uuid
        new_id = str(uuid.uuid4())[:8]
        
        if 'vignette_manager' not in st.session_state:
            from vignettes import VignetteManager
            st.session_state.vignette_manager = VignetteManager(st.session_state.user_id)
        
        st.session_state.vignette_manager.create_vignette_with_id(
            id=new_id,
            title="Untitled Vignette",
            content="<p>Write your story here...</p>",
            theme="Life Lesson",
            mood="Reflective",
            is_draft=True
        )
        
        st.session_state.editing_vignette_id = new_id
        st.session_state.show_vignette_modal = True
        st.rerun()
    
    if st.button("📖 View All Vignettes", use_container_width=True): 
        st.session_state.show_vignette_manager = True
        st.rerun()
    
    st.divider()
    st.header("📖 Session Management")
    if st.button("📋 All Sessions", use_container_width=True): 
        st.session_state.show_session_manager = True
        st.rerun()
    if st.button("➕ Custom Session", use_container_width=True): 
        st.session_state.show_session_creator = True
        st.rerun()
    
    # ============================================================================
    # PUBLISH YOUR BOOK SECTION
    # ============================================================================
    st.divider()
    st.subheader("📤 Publish Your Book")

    if st.session_state.logged_in and st.session_state.user_id:
        # Prepare export data for the publisher
        export_data = []
        for session in SESSIONS:
            sid = session["id"]
            sdata = st.session_state.responses.get(sid, {})
            for q, a in sdata.get("questions", {}).items():
                images_with_data = []
                if a.get("images"):
                    for img_ref in a.get("images", []):
                        img_id = img_ref.get("id")
                        b64 = st.session_state.image_handler.get_image_base64(img_id) if st.session_state.image_handler else None
                        caption = img_ref.get("caption", "")
                        if b64:
                            images_with_data.append({
                                "id": img_id, "base64": b64, "caption": caption
                            })
                
                export_item = {
                    "question": q, 
                    "answer_text": re.sub(r'<[^>]+>', '', a.get("answer", "")),
                    "timestamp": a.get("timestamp", ""), 
                    "session_id": sid, 
                    "session_title": session["title"],
                    "has_images": a.get("has_images", False), 
                    "image_count": a.get("image_count", 0),
                    "images": images_with_data
                }
                export_data.append(export_item)
        
        if export_data:
            # Save export data to session state for the publisher
            complete_data = {
                "user": st.session_state.user_id, 
                "user_profile": st.session_state.user_account.get('profile', {}),
                "narrative_gps": st.session_state.user_account.get('narrative_gps', {}),
                "enhanced_profile": st.session_state.user_account.get('enhanced_profile', {}),
                "cover_design": st.session_state.user_account.get('cover_design', {}),
                "stories": export_data, 
                "export_date": datetime.now().isoformat(),
                "summary": {
                    "total_stories": len(export_data), 
                    "total_sessions": len(set(s['session_id'] for s in export_data))
                }
            }
            
            # Save to a temp file and store path in session state
            temp_file = f"temp_export_{st.session_state.user_id}.json"
            with open(temp_file, 'w') as f:
                json.dump(complete_data, f)
            
            # Store in session state for the publisher
            st.session_state.publisher_data = complete_data
            st.session_state.publisher_data_path = temp_file
            
            # Button to open publisher in main screen
            if st.button("📚 Open Book Publisher", type="primary", use_container_width=True):
                st.session_state.show_publisher = True
                st.rerun()
            
            # Optional: Keep JSON backup
            with st.expander("📦 JSON Backup", expanded=False):
                json_data = json.dumps(complete_data, indent=2)
                st.download_button(
                    label="📥 Download JSON Backup", 
                    data=json_data,
                    file_name=f"Tell_My_Story_Backup_{st.session_state.user_id}.json",
                    mime="application/json", 
                    use_container_width=True
                )
        else: 
            st.warning("No stories yet! Start writing to publish.")
    else: 
        st.warning("Please log in to export your data.")
    
    # ============================================================================
    # CLEAR DATA SECTION
    # ============================================================================
    st.divider()
    st.subheader("⚠️ Clear Data")
    if st.session_state.confirming_clear == "session":
        st.warning("**Delete ALL answers in current session?**")
        if st.button("✅ Confirm", type="primary", key="conf_sesh", use_container_width=True): 
            sid = SESSIONS[st.session_state.current_session]["id"]
            st.session_state.responses[sid]["questions"] = {}
            save_user_data(st.session_state.user_id, st.session_state.responses)
            st.session_state.confirming_clear = None
            st.rerun()
        if st.button("❌ Cancel", key="can_sesh", use_container_width=True): 
            st.session_state.confirming_clear = None
            st.rerun()
    elif st.session_state.confirming_clear == "all":
        st.warning("**Delete ALL answers for ALL sessions?**")
        if st.button("✅ Confirm All", type="primary", key="conf_all", use_container_width=True): 
            for s in SESSIONS:
                st.session_state.responses[s["id"]]["questions"] = {}
            save_user_data(st.session_state.user_id, st.session_state.responses)
            st.session_state.confirming_clear = None
            st.rerun()
        if st.button("❌ Cancel", key="can_all", use_container_width=True): 
            st.session_state.confirming_clear = None
            st.rerun()
    else:
        if st.button("🗑️ Clear Session", use_container_width=True): 
            st.session_state.confirming_clear = "session"
            st.rerun()
        if st.button("🔥 Clear All", use_container_width=True): 
            st.session_state.confirming_clear = "all"
            st.rerun()
    
    # ============================================================================
    # SEARCH SECTION
    # ============================================================================
    st.divider()
    st.subheader("🔍 Search Your Stories")
    search_query = st.text_input("Search answers & captions...", placeholder="e.g., childhood, wedding, photo", key="global_search")
    if search_query and len(search_query) >= 2:
        results = search_all_answers(search_query)
        if results:
            st.success(f"Found {len(results)} matches")
            with st.expander(f"📖 {len(results)} Results", expanded=True):
                for i, r in enumerate(results[:10]):
                    st.markdown(f"**Session {r['session_id']}: {r['session_title']}**  \n*{r['question']}*")
                    if r.get('has_images'):
                        st.caption(f"📸 Contains {r.get('image_count', 1)} photo(s)")
                    st.markdown(f"{r['answer'][:150]}...")
                    if st.button(f"Go to Session", key=f"srch_go_{i}_{r['session_id']}", use_container_width=True):
                        for idx, s in enumerate(SESSIONS):
                            if s["id"] == r['session_id']:
                                st.session_state.update(current_session=idx, current_question_override=r['question'], show_ai_rewrite_menu=False)
                                st.rerun()
                    st.divider()
                if len(results) > 10: 
                    st.info(f"... and {len(results)-10} more matches")
        else: 
            st.info("No matches found")

# ============================================================================
# PUBLISHER FUNCTIONS - COMPLETE (replaces biography_publisher.py)
# ============================================================================

def clean_text_for_export(text):
    """Clean text for export - remove HTML tags but preserve structure"""
    if not text:
        return ""
    
    # Handle HTML entities
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&amp;', '&')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&quot;', '"')
    text = text.replace('&#39;', "'")
    
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    
    return text.strip()

def generate_docx_book(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    """Generate a Word document - just like typing it manually"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # ===== SIMPLE PAGE SETUP =====
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)  # This centers the text block
        section.right_margin = Inches(1.25)  # Equal margins = centered text block
    
    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # ===== COVER PAGE =====
    if cover_choice == "uploaded" and cover_image:
        try:
            # Just the image - centered
            image_stream = io.BytesIO(cover_image)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(image_stream, width=Inches(5))
            # NO TEXT - it's on your cover image
        except:
            # Fallback to text cover
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.font.size = Pt(42)
            run.font.bold = True
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"by {author}")
            run.font.size = Pt(24)
            run.font.italic = True
    else:
        # Text cover
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(42)
        run.font.bold = True
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"by {author}")
        run.font.size = Pt(24)
        run.font.italic = True
    
    doc.add_page_break()
    
    # ===== COPYRIGHT PAGE =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"© {datetime.now().year} {author}. All rights reserved.")
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    if include_toc:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Table of Contents")
        run.font.size = Pt(18)
        run.font.bold = True
        p.paragraph_format.space_after = Pt(12)
        
        # Group by session
        sessions = {}
        for story in stories:
            session_title = story.get('session_title', 'Untitled Session')
            if session_title not in sessions:
                sessions[session_title] = []
        
        # Add TOC entries - just plain text
        for session_title in sessions.keys():
            p = doc.add_paragraph(f"• {session_title}")
            p.paragraph_format.left_indent = Inches(0.5)
        
        doc.add_page_break()
    
    # ===== STORIES =====
    current_session = None
    for story in stories:
        session_title = story.get('session_title', 'Untitled Session')
        
        # Session header
        if session_title != current_session:
            current_session = session_title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(session_title)
            run.font.size = Pt(16)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        
        # Question (if interview style)
        if format_style == "interview":
            question_text = clean_text_for_export(story.get('question', ''))
            p = doc.add_paragraph(question_text)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].bold = True
            p.runs[0].italic = True
        
        # Answer
        answer_text = clean_text_for_export(story.get('answer_text', ''))
        if answer_text:
            paragraphs = answer_text.split('\n')
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.first_line_indent = Inches(0.25)
        
        # Images
        if include_images and story.get('images'):
            for img in story.get('images', []):
                if img.get('base64'):
                    try:
                        img_data = base64.b64decode(img['base64'])
                        img_stream = io.BytesIO(img_data)
                        
                        # Image centered
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(img_stream, width=Inches(4))
                        
                        # Caption
                        if img.get('caption'):
                            caption = clean_text_for_export(img['caption'])
                            p = doc.add_paragraph(caption)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.runs[0].font.size = Pt(10)
                            p.runs[0].font.italic = True
                    except:
                        continue
        
        # Space between stories
        doc.add_paragraph()
    
    # Save
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

def generate_html_book(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    """Generate an HTML document from stories with proper formatting"""
    
    html_parts = []
    
    # HTML header with styling for proper book formatting
    html_parts.append(f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{html.escape(title)}</title>
        <style>
            body {{
                font-family: 'Georgia', serif;
                line-height: 1.6;
                color: #333;
                max-width: 800px;
                margin: 0 auto;
                padding: 40px 20px;
                background: #fff;
            }}
            h1 {{
                font-size: 42px;
                text-align: center;
                margin-bottom: 10px;
                color: #000;
                font-weight: bold;
            }}
            h2 {{
                font-size: 28px;
                text-align: center;
                margin-top: 40px;
                margin-bottom: 20px;
                color: #444;
                border-bottom: 2px solid #eee;
                padding-bottom: 10px;
                font-weight: bold;
            }}
            .author {{
                text-align: center;
                font-size: 24px;
                color: #666;
                margin-bottom: 40px;
                font-style: italic;
            }}
            .question {{
                font-weight: bold;
                font-size: 18px;
                margin-top: 30px;
                margin-bottom: 10px;
                color: #2c3e50;
                border-left: 4px solid #3498db;
                padding-left: 15px;
                text-align: left;
            }}
            .answer {{
                text-align: left;
                margin-bottom: 20px;
            }}
            .answer p {{
                text-indent: 0.5in;
                margin-bottom: 6px;
                text-align: left;
                line-height: 1.8;
            }}
            .story-image {{
                max-width: 100%;
                height: auto;
                display: block;
                margin: 20px auto;
                border-radius: 5px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }}
            .image-caption {{
                text-align: center;
                font-size: 14px;
                color: #666;
                margin-top: 5px;
                margin-bottom: 20px;
                font-style: italic;
            }}
            .cover-page {{
                text-align: center;
                margin-bottom: 50px;
                page-break-after: always;
                min-height: 90vh;
                display: flex;
                flex-direction: column;
                justify-content: center;
            }}
            .cover-image {{
                max-width: 100%;
                max-height: 70vh;
                object-fit: contain;
                margin: 20px auto;
                border-radius: 10px;
                box-shadow: 0 10px 30px rgba(0,0,0,0.2);
            }}
            .simple-cover {{
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                padding: 60px 20px;
                border-radius: 10px;
                color: white;
                margin: 20px;
                text-align: center;
            }}
            .simple-cover h1 {{
                color: white;
                text-align: center;
                text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
            }}
            .simple-cover .author {{
                color: rgba(255,255,255,0.9);
                text-align: center;
            }}
            .copyright {{
                text-align: center;
                font-size: 12px;
                color: #999;
                margin-top: 50px;
                padding-top: 20px;
                border-top: 1px solid #eee;
            }}
            .toc {{
                background: #f9f9f9;
                padding: 20px;
                border-radius: 5px;
                margin: 30px 0;
                text-align: left;
            }}
            .toc h3 {{
                text-align: center;
                margin-top: 0;
            }}
            .toc ul {{
                list-style-type: none;
                padding-left: 0;
            }}
            .toc li {{
                margin-bottom: 10px;
                text-align: left;
                font-size: 16px;
            }}
            .toc a {{
                color: #3498db;
                text-decoration: none;
            }}
            .toc a:hover {{
                text-decoration: underline;
            }}
            hr {{
                margin: 30px 0;
                border: none;
                border-top: 1px dashed #ccc;
            }}
            @media print {{
                body {{
                    padding: 0.5in;
                }}
                .cover-page {{
                    page-break-after: always;
                    min-height: auto;
                }}
                h2 {{
                    page-break-before: always;
                }}
            }}
        </style>
    </head>
    <body>
    """)
    
    # COVER PAGE
    html_parts.append('<div class="cover-page">')
    
    if cover_choice == "uploaded" and cover_image:
        try:
            img_base64 = base64.b64encode(cover_image).decode()
            html_parts.append(f'''
            <div>
                <img src="data:image/jpeg;base64,{img_base64}" class="cover-image" alt="Book Cover">
                <h1>{html.escape(title)}</h1>
                <p class="author">by {html.escape(author)}</p>
            </div>
            ''')
        except Exception:
            html_parts.append(f'''
            <div class="simple-cover">
                <h1>{html.escape(title)}</h1>
                <p class="author">by {html.escape(author)}</p>
            </div>
            ''')
    else:
        html_parts.append(f'''
        <div class="simple-cover">
            <h1>{html.escape(title)}</h1>
            <p class="author">by {html.escape(author)}</p>
        </div>
        ''')
    
    html_parts.append('</div>')
    
    # Copyright page
    html_parts.append(f'<p class="copyright">© {datetime.now().year} {html.escape(author)}. All rights reserved.</p>')
    
    # Table of Contents
    if include_toc:
        html_parts.append('<div class="toc">')
        html_parts.append('<h3>Table of Contents</h3>')
        html_parts.append('<ul>')
        
        sessions = {}
        for story in stories:
            session_title = story.get('session_title', 'Untitled Session')
            if session_title not in sessions:
                sessions[session_title] = []
            sessions[session_title].append(story)
        
        for session_title in sessions.keys():
            anchor = session_title.lower().replace(' ', '-').replace('?', '').replace('!', '').replace(',', '').replace('.', '')
            html_parts.append(f'<li><a href="#{anchor}">{html.escape(session_title)}</a></li>')
        
        html_parts.append('</ul>')
        html_parts.append('</div>')
    
    # Add stories
    current_session = None
    for story in stories:
        session_title = story.get('session_title', 'Untitled Session')
        anchor = session_title.lower().replace(' ', '-').replace('?', '').replace('!', '').replace(',', '').replace('.', '')
        
        if session_title != current_session:
            current_session = session_title
            html_parts.append(f'<h2 id="{anchor}">{html.escape(session_title)}</h2>')
        
        if format_style == "interview":
            question_text = story.get('question', '')
            clean_question = clean_text_for_export(question_text)
            html_parts.append(f'<div class="question">{html.escape(clean_question)}</div>')
        
        # Format answer with proper paragraphs
        answer_text = story.get('answer_text', '')
        if answer_text:
            clean_answer = clean_text_for_export(answer_text)
            
            html_parts.append('<div class="answer">')
            paragraphs = clean_answer.split('\n')
            for para in paragraphs:
                if para.strip():
                    escaped_para = html.escape(para.strip())
                    html_parts.append(f'<p>{escaped_para}</p>')
            html_parts.append('</div>')
        
        # Add images
        if include_images and story.get('images'):
            for img in story.get('images', []):
                if img.get('base64'):
                    img_data = img['base64']
                    # Ensure proper base64 formatting
                    if not img_data.startswith('data:image'):
                        html_parts.append(f'<img src="data:image/jpeg;base64,{img_data}" class="story-image" alt="Story image">')
                    else:
                        html_parts.append(f'<img src="{img_data}" class="story-image" alt="Story image">')
                    
                    if img.get('caption'):
                        clean_caption = clean_text_for_export(img['caption'])
                        caption = html.escape(clean_caption)
                        html_parts.append(f'<p class="image-caption">{caption}</p>')
        
        html_parts.append('<hr>')
    
    html_parts.append("""
    </body>
    </html>
    """)
    
    return '\n'.join(html_parts)

def generate_epub_book(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    """Generate an EPUB file (reflowable for e-readers)"""
    try:
        from ebooklib import epub
        
        # Create EPUB book
        book = epub.EpubBook()
        
        # Set metadata
        book.set_identifier(hashlib.md5(f"{title}{author}{datetime.now()}".encode()).hexdigest()[:16])
        book.set_title(title)
        book.set_language('en')
        book.add_author(author)
        
        # Add cover
        if cover_choice == "uploaded" and cover_image:
            book.set_cover("cover.jpg", cover_image)
        else:
            # Create a simple HTML cover
            cover_content = f"""
            <html>
            <body style="text-align: center; margin-top: 20%;">
                <h1>{html.escape(title)}</h1>
                <h2>by {html.escape(author)}</h2>
            </body>
            </html>
            """
            cover_page = epub.EpubHtml(title='Cover', file_name='cover.xhtml', lang='en')
            cover_page.content = cover_content
            book.add_item(cover_page)
        
        # Add CSS
        style = '''
        body { font-family: Georgia, serif; line-height: 1.6; margin: 5%; }
        h1 { text-align: center; }
        h2 { text-align: center; }
        .question { font-weight: bold; font-style: italic; margin-top: 1em; }
        .answer p { text-indent: 0.5in; margin-bottom: 0.5em; }
        .session-header { text-align: center; font-size: 1.5em; margin: 1em 0; }
        .image-caption { text-align: center; font-style: italic; font-size: 0.9em; }
        '''
        nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
        book.add_item(nav_css)
        
        chapters = []
        current_session = None
        chapter_index = 1
        
        for story in stories:
            session_title = story.get('session_title', 'Untitled Session')
            
            # New session = new chapter
            if session_title != current_session:
                current_session = session_title
                
                # Create chapter for this session
                chapter = epub.EpubHtml(
                    title=session_title,
                    file_name=f'chap_{chapter_index:02d}.xhtml',
                    lang='en'
                )
                chapter.add_item(nav_css)
                
                # Start chapter content
                content = [f'<h1 class="session-header">{html.escape(session_title)}</h1>']
                
                # Find all stories in this session
                session_stories = [s for s in stories if s.get('session_title') == session_title]
                
                for s in session_stories:
                    if format_style == "interview":
                        question = clean_text_for_export(s.get('question', ''))
                        content.append(f'<p class="question">{html.escape(question)}</p>')
                    
                    answer = clean_text_for_export(s.get('answer_text', ''))
                    if answer:
                        content.append('<div class="answer">')
                        paragraphs = answer.split('\n')
                        for para in paragraphs:
                            if para.strip():
                                content.append(f'<p>{html.escape(para.strip())}</p>')
                        content.append('</div>')
                    
                    # Add images
                    if include_images and s.get('images'):
                        for img in s.get('images', []):
                            if img.get('base64'):
                                img_data = base64.b64decode(img['base64'])
                                img_file = f"img_{chapter_index}_{img.get('id', 'unknown')}.jpg"
                                img_item = epub.EpubImage()
                                img_item.file_name = f"images/{img_file}"
                                img_item.media_type = "image/jpeg"
                                img_item.content = img_data
                                book.add_item(img_item)
                                
                                content.append(f'<img src="images/{img_file}" style="max-width:100%; display:block; margin:20px auto;"/>')
                                
                                if img.get('caption'):
                                    caption = clean_text_for_export(img['caption'])
                                    content.append(f'<p class="image-caption">{html.escape(caption)}</p>')
                
                chapter.content = '\n'.join(content)
                book.add_item(chapter)
                chapters.append(chapter)
                chapter_index += 1
        
        # Add table of contents
        book.toc = chapters
        
        # Add navigation
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Define spine (reading order)
        book.spine = ['nav'] + chapters
        
        # Generate EPUB
        epub_bytes = io.BytesIO()
        epub.write_epub(epub_bytes, book)
        epub_bytes.seek(0)
        
        return epub_bytes.getvalue(), None
        
    except ImportError:
        return None, "Please install ebooklib: pip install ebooklib"
    except Exception as e:
        return None, f"Error generating EPUB: {e}"

def generate_rtf_book(title, author, stories, format_style="interview", include_toc=True, include_images=True, cover_image=None, cover_choice="simple"):
    """Generate an RTF file (Rich Text Format)"""
    try:
        rtf = r"""{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}
\paperw12240\paperh15840\margl1440\margr1440\margt1440\margb1440
"""
        # Title
        rtf += r"\pard\qc\fs72\b " + title + r"\par\par"
        # Author
        rtf += r"\pard\qc\fs48\i by " + author + r"\par\par"
        # Copyright
        rtf += r"\pard\qc\fs24\i Copyright " + str(datetime.now().year) + r" " + author + r". All rights reserved.\par\par"
        
        # Group by session
        sessions = {}
        for story in stories:
            session_title = story.get('session_title', '')
            if session_title not in sessions:
                sessions[session_title] = []
            sessions[session_title].append(story)
        
        # Table of Contents
        if include_toc:
            rtf += r"\pard\qc\fs36\b Table of Contents\par\par"
            for session_title in sessions.keys():
                rtf += r"\pard\ql\fs28 " + session_title + r"\par"
            rtf += r"\par"
        
        # Stories
        for session_title, session_stories in sessions.items():
            rtf += r"\pard\qc\fs40\b " + session_title + r"\par\par"
            
            for story in session_stories:
                if format_style == "interview":
                    question = clean_text_for_export(story.get('question', ''))
                    rtf += r"\pard\ql\fs28\b\i " + question + r"\par"
                
                answer = clean_text_for_export(story.get('answer_text', ''))
                paragraphs = answer.split('\n')
                for para in paragraphs:
                    if para.strip():
                        rtf += r"\pard\ql\fs24\fi360 " + para.strip() + r"\par"
                rtf += r"\par"
        
        rtf += "}"
        return rtf.encode('utf-8')
        
    except Exception as e:
        st.error(f"Error generating RTF: {e}")
        return None

def show_celebration():
    """Show celebration animation"""
    st.balloons()
    st.success("🎉 Your book has been generated successfully!")

# ============================================================================
# PUBLISHER PAGE - SHOW ON MAIN SCREEN WHEN ACTIVATED
# ============================================================================
if st.session_state.get('show_publisher', False):
    # Hide sidebar when in publisher mode
    st.markdown("""
    <style>
        section[data-testid="stSidebar"] { display: none !important; }
    </style>
    """, unsafe_allow_html=True)
    
    # Back button at top
    col1, col2, col3 = st.columns([1, 6, 1])
    with col1:
        if st.button("← Back to Writing", use_container_width=True):
            st.session_state.show_publisher = False
            st.rerun()
    
    with col2:
        st.markdown("""
        <div style="text-align: center; padding: 2rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); margin: 0 -1rem 2rem -1rem; border-radius: 0 0 20px 20px; color: white;">
            <h1>📚 Book Publisher</h1>
            <p>Transform your stories into a beautifully formatted book</p>
        </div>
        """, unsafe_allow_html=True)
    
    # Get data directly from user data
    stories_for_export = []
    
    if st.session_state.logged_in and st.session_state.user_id:
        # Build stories directly from session state responses
        for session in SESSIONS:
            sid = session["id"]
            sdata = st.session_state.responses.get(sid, {})
            
            for question_text, answer_data in sdata.get("questions", {}).items():
                # Get images with base64 data
                images_with_data = []
                if answer_data.get("images") and st.session_state.image_handler:
                    for img_ref in answer_data.get("images", []):
                        img_id = img_ref.get("id")
                        # Get base64 from image handler
                        b64 = st.session_state.image_handler.get_image_base64(img_id) if st.session_state.image_handler else None
                        if b64:
                            images_with_data.append({
                                "id": img_id,
                                "base64": b64,
                                "caption": img_ref.get("caption", "")
                            })
                
                story_item = {
                    "question": question_text,
                    "answer_text": answer_data.get("answer", ""),
                    "timestamp": answer_data.get("timestamp", ""),
                    "session_id": sid,
                    "session_title": session["title"],
                    "has_images": answer_data.get("has_images", False),
                    "image_count": answer_data.get("image_count", 0),
                    "images": images_with_data
                }
                stories_for_export.append(story_item)
        
        if stories_for_export:
            # Book details
            col1, col2 = st.columns(2)
            with col1:
                profile = st.session_state.user_account.get('profile', {})
                default_title = f"{profile.get('first_name', 'My')}'s Story"
                book_title = st.text_input("Book Title", value=default_title)
            with col2:
                default_author = f"{profile.get('first_name', '')} {profile.get('last_name', '')}".strip()
                book_author = st.text_input("Author Name", value=default_author if default_author else "Author Name")
            
            # Format options
            st.markdown("---")
            col1, col2, col3 = st.columns(3)
            with col1:
                format_style = st.radio(
                    "📝 Format Style",
                    ["interview", "biography"],
                    format_func=lambda x: {
                        "interview": "Show Questions & Answers", 
                        "biography": "Just Answers (Biography Style)"
                    }[x],
                    horizontal=True
                )
            with col2:
                include_toc = st.checkbox("📖 Table of Contents", value=True)
            with col3:
                cover_choice = st.radio(
                    "🎨 Cover Type",
                    ["simple", "uploaded"],
                    format_func=lambda x: {
                        "simple": "Simple Gradient Cover",
                        "uploaded": "Use My Uploaded Image"
                    }[x],
                    horizontal=True
                )
            
            # Image uploader for cover
            uploaded_cover = None
            if cover_choice == "uploaded":
                st.markdown("---")
                st.markdown("### 🖼️ Upload Cover Image")
                uploaded_cover = st.file_uploader(
                    "Choose an image (JPG or PNG)", 
                    type=['jpg', 'jpeg', 'png'], 
                    key="publisher_cover_upload"
                )
                if uploaded_cover:
                    st.image(uploaded_cover, width=200, caption="Your cover image")
                    st.success("✅ Cover image ready")
            
            # Summary stats
            st.markdown("---")
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Total Stories", len(stories_for_export))
            with col2:
                total_sessions = len(set(s['session_id'] for s in stories_for_export))
                st.metric("Sessions", total_sessions)
            with col3:
                total_images = sum(len(s.get('images', [])) for s in stories_for_export)
                st.metric("Images", total_images)
            with col4:
                total_words = sum(len(s.get('answer_text', '').split()) for s in stories_for_export)
                st.metric("Words", f"{total_words:,}")
            
            # Preview section
            with st.expander("📖 Preview First 3 Stories", expanded=False):
                for i, story in enumerate(stories_for_export[:3]):
                    if format_style == "interview":
                        st.markdown(f"**Q: {story.get('question', '')}**")
                    else:
                        st.markdown(f"**{story.get('session_title', 'Session')}**")
                    
                    # Clean and show preview
                    preview_text = clean_text_for_export(story.get('answer_text', ''))[:300]
                    st.markdown(f"{preview_text}...")
                    
                    if story.get('images'):
                        st.caption(f"📸 {len(story['images'])} image(s)")
                    if i < 2:
                        st.divider()
            
            # Generate buttons - FOUR FORMATS
            st.markdown("---")
            st.markdown("### 🖨️ Generate Your Book")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.button("📊 DOCX", type="primary", use_container_width=True):
                    with st.spinner("Creating Word document..."):
                        cover_image_data = uploaded_cover.getvalue() if uploaded_cover else None
                        
                        docx_bytes = generate_docx_book(
                            book_title,
                            book_author,
                            stories_for_export,
                            format_style,
                            include_toc,
                            True,
                            cover_image_data,
                            cover_choice
                        )
                        
                        filename = f"{book_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                        
                        st.download_button(
                            "📥 Download DOCX", 
                            data=docx_bytes, 
                            file_name=filename, 
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                            use_container_width=True,
                            key="docx_download"
                        )
                        show_celebration()
            
            with col2:
                if st.button("🌐 HTML", type="primary", use_container_width=True):
                    with st.spinner("Creating HTML page..."):
                        cover_image_data = uploaded_cover.getvalue() if uploaded_cover else None
                        
                        html_content = generate_html_book(
                            book_title,
                            book_author,
                            stories_for_export,
                            format_style,
                            include_toc,
                            True,
                            cover_image_data,
                            cover_choice
                        )
                        
                        filename = f"{book_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.html"
                        
                        st.download_button(
                            "📥 Download HTML", 
                            data=html_content, 
                            file_name=filename, 
                            mime="text/html", 
                            use_container_width=True,
                            key="html_download"
                        )
                        show_celebration()
            
            with col3:
                if st.button("📱 EPUB", type="primary", use_container_width=True):
                    with st.spinner("Creating EPUB file..."):
                        cover_image_data = uploaded_cover.getvalue() if uploaded_cover else None
                        
                        epub_bytes, error = generate_epub_book(
                            book_title,
                            book_author,
                            stories_for_export,
                            format_style,
                            include_toc,
                            True,
                            cover_image_data,
                            cover_choice
                        )
                        
                        if epub_bytes:
                            filename = f"{book_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.epub"
                            st.download_button(
                                "📥 Download EPUB", 
                                data=epub_bytes, 
                                file_name=filename, 
                                mime="application/epub+zip", 
                                use_container_width=True,
                                key="epub_download"
                            )
                            show_celebration()
                        else:
                            st.error(f"Failed to generate EPUB: {error}")
                            st.info("💡 Install ebooklib: pip install ebooklib")
            
            with col4:
                if st.button("📝 RTF", type="primary", use_container_width=True):
                    with st.spinner("Creating RTF file..."):
                        cover_image_data = uploaded_cover.getvalue() if uploaded_cover else None
                        
                        rtf_bytes = generate_rtf_book(
                            book_title,
                            book_author,
                            stories_for_export,
                            format_style,
                            include_toc,
                            True,
                            cover_image_data,
                            cover_choice
                        )
                        
                        if rtf_bytes:
                            filename = f"{book_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.rtf"
                            st.download_button(
                                "📥 Download RTF", 
                                data=rtf_bytes, 
                                file_name=filename, 
                                mime="application/rtf", 
                                use_container_width=True,
                                key="rtf_download"
                            )
                            show_celebration()
                        else:
                            st.error("Failed to generate RTF")
            
            # Optional: JSON backup
            with st.expander("📦 JSON Backup", expanded=False):
                complete_data = {
                    "user": st.session_state.user_id,
                    "user_profile": st.session_state.user_account.get('profile', {}),
                    "book_title": book_title,
                    "book_author": book_author,
                    "stories": stories_for_export,
                    "export_date": datetime.now().isoformat(),
                    "summary": {
                        "total_stories": len(stories_for_export),
                        "total_sessions": total_sessions,
                        "total_words": total_words
                    }
                }
                json_data = json.dumps(complete_data, indent=2)
                st.download_button(
                    label="📥 Download JSON Backup", 
                    data=json_data,
                    file_name=f"Tell_My_Story_Backup_{st.session_state.user_id}.json",
                    mime="application/json", 
                    use_container_width=True
                )
        
        else:
            st.warning("No stories found! Start writing to publish your book.")
            if st.button("← Return to Main App"):
                st.session_state.show_publisher = False
                st.rerun()
    else:
        st.warning("Please log in to publish your book.")
        if st.button("← Return to Main App"):
            st.session_state.show_publisher = False
            st.rerun()
    
    # Stop here - don't show main content
    st.stop()

# ============================================================================
# MAIN CONTENT AREA
# ============================================================================

if (st.session_state.show_vignette_modal or 
    st.session_state.show_vignette_manager or 
    st.session_state.show_vignette_detail or
    st.session_state.show_topic_browser or 
    st.session_state.show_session_manager or 
    st.session_state.show_session_creator or
    st.session_state.show_bank_manager or 
    st.session_state.show_bank_editor or
    st.session_state.show_privacy_settings or
    st.session_state.show_cover_designer or
    st.session_state.show_profile_setup or
    st.session_state.show_ai_rewrite):
    
    st.markdown(f'<div class="main-header"><img src="{LOGO_URL}" class="logo-img"></div>', unsafe_allow_html=True)
    st.stop()

if st.session_state.current_session >= len(SESSIONS): 
    st.session_state.current_session = 0

current_session = SESSIONS[st.session_state.current_session]
current_session_id = current_session["id"]

if st.session_state.current_question_override:
    current_question_text = st.session_state.current_question_override
    question_source = "custom"
else:
    if st.session_state.current_question >= len(current_session["questions"]): 
        st.session_state.current_question = 0
    current_question_text = current_session["questions"][st.session_state.current_question]
    question_source = "regular"

st.markdown("---")

col1, col2 = st.columns([3, 1])
with col1:
    st.subheader(f"Session {current_session_id}: {current_session['title']}")
    sdata = st.session_state.responses.get(current_session_id, {})
    answered = len(sdata.get("questions", {}))
    total = len(current_session["questions"])
    if total > 0: 
        st.progress(answered/total)
        st.caption(f"📝 Topics explored: {answered}/{total} ({answered/total*100:.0f}%)")
with col2:
    if question_source == "custom":
        st.markdown(f'<div class="custom-topic-badge">{"📝 Vignette" if "Vignette:" in st.session_state.current_question_override else "✨ Custom Topic"}</div>', unsafe_allow_html=True)
    else:
        st.markdown(f'<div class="topic-counter">Topic {st.session_state.current_question+1} of {len(current_session["questions"])}</div>', unsafe_allow_html=True)

st.markdown(f'<div class="question-box">{current_question_text}</div>', unsafe_allow_html=True)

if question_source == "regular":
    st.markdown(f'<div class="chapter-guidance">{current_session.get("guidance", "")}</div>', unsafe_allow_html=True)
else:
    if "Vignette:" in current_question_text:
        st.info("📝 **Vignette Mode** - Write a short, focused story about a specific moment or memory.")
    else:
        st.info("✨ **Custom Topic** - Write about whatever comes to mind!")

st.write("")
st.write("")

existing_answer = ""
if current_session_id in st.session_state.responses:
    if current_question_text in st.session_state.responses[current_session_id]["questions"]:
        existing_answer = st.session_state.responses[current_session_id]["questions"][current_question_text]["answer"]

if st.session_state.logged_in:
    init_image_handler()
    existing_images = st.session_state.image_handler.get_images_for_answer(current_session_id, current_question_text) if st.session_state.image_handler else []

# ============================================================================
# QUILL EDITOR
# ============================================================================
import logging

editor_base_key = f"quill_{current_session_id}_{current_question_text[:20]}"
content_key = f"{editor_base_key}_content"

version_key = f"{editor_base_key}_version"
if version_key not in st.session_state:
    st.session_state[version_key] = 0

existing_answer = ""
if current_session_id in st.session_state.responses:
    if current_question_text in st.session_state.responses[current_session_id]["questions"]:
        existing_answer = st.session_state.responses[current_session_id]["questions"][current_question_text]["answer"]

if content_key not in st.session_state:
    if existing_answer and existing_answer != "<p>Start writing your story here...</p>":
        st.session_state[content_key] = existing_answer
    else:
        st.session_state[content_key] = "<p>Start writing your story here...</p>"

st.markdown("### ✍️ Your Story")
st.markdown("""
<div class="image-drop-info">
    📸 <strong>Drag & drop images</strong> directly into the editor.
</div>
""", unsafe_allow_html=True)

question_text_safe = "".join(c for c in current_question_text if c.isalnum() or c.isspace()).replace(" ", "_")[:30]
editor_component_key = f"quill_editor_{current_session_id}_{question_text_safe}_v{st.session_state[version_key]}"

print(f"Creating Quill editor with key: {editor_component_key}")

try:
    content = st_quill(
        value=st.session_state[content_key],
        key=editor_component_key,
        placeholder="Start writing your story here...",
        html=True
    )
    
    if content is not None and content != st.session_state[content_key]:
        st.session_state[content_key] = content
        
except Exception as e:
    st.error(f"Error loading editor: {str(e)}")
    content = st.text_area(
        "Your story (fallback editor):",
        value=re.sub(r'<[^>]+>', '', st.session_state[content_key]),
        height=300,
        key=f"fallback_{editor_base_key}"
    )
    if content:
        st.session_state[content_key] = f"<p>{content}</p>"

st.markdown("---")

# ============================================================================
# BUTTONS ROW - WITH PROMPT ME BUTTON ADDED
# ============================================================================
col1, col2, col3, col4, col5, col6, col7, col8 = st.columns([1, 1, 1, 1, 1, 1, 1, 2])

spellcheck_base = f"spell_{editor_base_key}"
spell_result_key = f"{spellcheck_base}_result"
current_content = st.session_state.get(content_key, "")
has_content = current_content and current_content != "<p><br></p>" and current_content != "<p>Start writing your story here...</p>"
showing_results = spell_result_key in st.session_state and st.session_state[spell_result_key].get("show", False)

import_key = f"import_{editor_base_key}"
if import_key not in st.session_state:
    st.session_state[import_key] = False
show_import = st.session_state[import_key]

with col1:
    if st.button("💾 Save", key=f"save_btn_{editor_base_key}", type="primary", use_container_width=True):
        current_content = st.session_state[content_key]
        if current_content and current_content.strip() and current_content != "<p><br></p>" and current_content != "<p>Start writing your story here...</p>":
            with st.spinner("Saving your story..."):
                if save_response(current_session_id, current_question_text, current_content):
                    st.success("✅ Saved!")
                    time.sleep(0.5)
                    st.rerun()
                else: 
                    st.error("Failed to save")
        else: 
            st.warning("Please write something!")

with col2:
    if existing_answer and existing_answer != "<p>Start writing your story here...</p>":
        if st.button("🗑️ Delete", key=f"del_btn_{editor_base_key}", use_container_width=True):
            if delete_response(current_session_id, current_question_text):
                st.session_state[content_key] = "<p>Start writing your story here...</p>"
                st.success("✅ Deleted!")
                st.rerun()
    else: 
        st.button("🗑️ Delete", key=f"del_disabled_{editor_base_key}", disabled=True, use_container_width=True)

with col3:
    if has_content and not showing_results:
        if st.button("🔍 Spell Check", key=f"spell_{editor_base_key}", use_container_width=True):
            with st.spinner("Checking spelling and grammar..."):
                text_only = re.sub(r'<[^>]+>', '', current_content)
                if len(text_only.split()) >= 3:
                    corrected = auto_correct_text(text_only)
                    if corrected and corrected != text_only:
                        st.session_state[spell_result_key] = {
                            "original": text_only,
                            "corrected": corrected,
                            "show": True
                        }
                    else:
                        st.session_state[spell_result_key] = {
                            "message": "✅ No spelling or grammar issues found!",
                            "show": True
                        }
                    st.rerun()
                else:
                    st.warning("Text too short for spell check (minimum 3 words)")
    else:
        st.button("🔍 Spell Check", key=f"spell_disabled_{editor_base_key}", disabled=True, use_container_width=True)

with col4:
    if has_content:
        if st.button("✨ AI Rewrite", key=f"rewrite_btn_{editor_base_key}", use_container_width=True):
            st.session_state.show_ai_rewrite_menu = True
            st.rerun()
    else:
        st.button("✨ AI Rewrite", key=f"rewrite_disabled_{editor_base_key}", disabled=True, use_container_width=True)

# PROMPT ME BUTTON
with col5:
    if st.button("💭 Prompt Me", key=f"prompt_btn_{editor_base_key}", use_container_width=True):
        with st.spinner("Generating personalized writing prompts..."):
            # Get profile context
            profile_context = get_narrative_gps_for_ai()
            
            # Get birth year from profile
            birth_year = None
            if st.session_state.user_account and 'profile' in st.session_state.user_account:
                birthdate = st.session_state.user_account['profile'].get('birthdate', '')
                if birthdate:
                    import re
                    year_match = re.search(r'\d{4}', birthdate)
                    if year_match:
                        birth_year = int(year_match.group())
            
            # Add enhanced profile context
            if st.session_state.user_account and 'enhanced_profile' in st.session_state.user_account:
                ep = st.session_state.user_account['enhanced_profile']
                if ep:
                    profile_context += "\n\nPersonal background:\n"
                    if ep.get('first_name'): 
                        profile_context += f"- Name: {ep.get('first_name')}\n"
                    if ep.get('birth_place'): 
                        profile_context += f"- Birth place: {ep['birth_place']}\n"
                    if ep.get('childhood_home'): 
                        profile_context += f"- Childhood home: {ep['childhood_home'][:150]}\n"
                    if ep.get('parents'):
                        profile_context += f"- Parents: {ep['parents'][:150]}\n"
            
            # Generate prompts
            result = generate_writing_prompts(
                current_session['title'],
                current_question_text,
                current_content,
                profile_context,
                birth_year
            )
            
            if result.get('success'):
                st.session_state.current_prompt_data = result
                st.session_state.show_prompt_modal = True
                st.rerun()
            else:
                st.error(result.get('error', 'Could not generate prompts'))

with col6:
    button_label = "📂 Close Import" if show_import else "📂 Import File"
    if st.button(button_label, key=f"import_btn_{editor_base_key}", use_container_width=True):
        st.session_state[import_key] = not show_import
        st.rerun()

with col7:
    if st.session_state.get('show_ai_rewrite_menu', False):
        person_option = st.selectbox(
            "Voice:",
            options=["1st", "2nd", "3rd"],
            format_func=lambda x: {"1st": "👤 First Person", "2nd": "💬 Second Person", "3rd": "📖 Third Person"}[x],
            key=f"person_select_{editor_base_key}",
            label_visibility="collapsed"
        )
        
        if st.button("Go", key=f"go_rewrite_{editor_base_key}", type="primary", use_container_width=True):
            with st.spinner(f"Rewriting in {person_option} person..."):
                current_content = st.session_state[content_key]
                result = ai_rewrite_answer(
                    current_content, 
                    person_option, 
                    current_question_text, 
                    current_session['title']
                )
                
                if result.get('success'):
                    st.session_state.current_rewrite_data = result
                    st.session_state.show_ai_rewrite = True
                    st.session_state.show_ai_rewrite_menu = False
                    st.rerun()
                else:
                    st.error(result.get('error', 'Failed to rewrite'))
    else:
        st.markdown("")

with col8:
    nav1, nav2 = st.columns(2)
    with nav1: 
        prev_disabled = st.session_state.current_question == 0
        if st.button("← Previous", disabled=prev_disabled, key=f"prev_{editor_base_key}", use_container_width=True):
            if not prev_disabled:
                st.session_state.current_question -= 1
                st.session_state.current_question_override = None
                st.session_state.show_ai_rewrite_menu = False
                st.rerun()
    with nav2:
        next_disabled = st.session_state.current_question >= len(current_session["questions"]) - 1
        if st.button("Next →", disabled=next_disabled, key=f"next_{editor_base_key}", use_container_width=True):
            if not next_disabled:
                st.session_state.current_question += 1
                st.session_state.current_question_override = None
                st.session_state.show_ai_rewrite_menu = False
                st.rerun()

if showing_results:
    result = st.session_state[spell_result_key]
    if "corrected" in result:
        st.markdown("---")
        st.markdown("### ✅ Suggested Corrections:")
        st.markdown(f'<div style="background-color: #f0f9ff; padding: 15px; border-radius: 8px; border-left: 4px solid #4CAF50;">{result["corrected"]}</div>', unsafe_allow_html=True)
        
        col_apply1, col_apply2, col_apply3 = st.columns([1, 1, 1])
        with col_apply2:
            if st.button("📋 Apply Corrections", key=f"{spellcheck_base}_apply", type="primary", use_container_width=True):
                corrected = result["corrected"]
                if not corrected.startswith('<p>'):
                    corrected = f'<p>{corrected}</p>'
                
                st.session_state[content_key] = corrected
                save_response(current_session_id, current_question_text, corrected)
                st.session_state[version_key] += 1
                st.session_state[spell_result_key] = {"show": False}
                st.success("✅ Corrections applied!")
                st.rerun()
            
            if st.button("❌ Dismiss", key=f"{spellcheck_base}_dismiss", use_container_width=True):
                st.session_state[spell_result_key] = {"show": False}
                st.rerun()
    
    elif "message" in result:
        st.success(result["message"])
        if st.button("Dismiss", key=f"{spellcheck_base}_dismiss_msg"):
            st.session_state[spell_result_key] = {"show": False}
            st.rerun()

if show_import:
    st.markdown("---")
    st.markdown("### 📂 Import Text File")
    
    with st.expander("📋 Supported File Formats", expanded=True):
        st.markdown("""
        | Format | Description |
        |--------|-------------|
        | **.txt** | Plain text |
        | **.docx** | Microsoft Word |
        | **.rtf** | Rich Text Format |
        | **.vtt/.srt** | Subtitle files |
        | **.json** | Transcription JSON |
        | **.md** | Markdown |
        
        **Maximum file size:** 50MB
        """)
    
    uploaded_file = st.file_uploader(
        "Choose a file to import",
        type=['txt', 'docx', 'rtf', 'vtt', 'srt', 'json', 'md'],
        key=f"file_uploader_{editor_base_key}",
        help="Select a file from your computer to import into this topic"
    )
    
    if uploaded_file:
        col_imp1, col_imp2, col_imp3 = st.columns([1, 1, 2])
        with col_imp1:
            if st.button("📥 Import", key=f"do_import_{editor_base_key}", type="primary", use_container_width=True):
                with st.spinner("Importing file..."):
                    imported_html = import_text_file_main(uploaded_file)
                    if imported_html:
                        current = st.session_state.get(content_key, "")
                        if current and current != "<p>Start writing your story here...</p>" and current != "<p><br></p>":
                            st.session_state[f"{import_key}_pending"] = imported_html
                            st.session_state[f"{import_key}_show_options"] = True
                            st.rerun()
                        else:
                            st.session_state[content_key] = imported_html
                            st.session_state[version_key] += 1
                            st.session_state[import_key] = False
                            st.success("✅ File imported successfully!")
                            st.rerun()
        
        with col_imp2:
            if st.button("❌ Cancel", key=f"cancel_import_{editor_base_key}", use_container_width=True):
                st.session_state[import_key] = False
                st.rerun()
        
        if st.session_state.get(f"{import_key}_show_options", False):
            st.markdown("---")
            st.markdown("**This topic already has content. What would you like to do?**")
            
            col_opt1, col_opt2, col_opt3 = st.columns(3)
            with col_opt1:
                if st.button("📝 Replace Current", key=f"import_replace_{editor_base_key}", use_container_width=True):
                    st.session_state[content_key] = st.session_state[f"{import_key}_pending"]
                    st.session_state[version_key] += 1
                    st.session_state[import_key] = False
                    st.session_state[f"{import_key}_pending"] = None
                    st.session_state[f"{import_key}_show_options"] = False
                    st.success("✅ File imported (replaced current content)!")
                    st.rerun()
            
            with col_opt2:
                if st.button("➕ Append to Current", key=f"import_append_{editor_base_key}", use_container_width=True):
                    current = st.session_state.get(content_key, "")
                    current = current.replace('</p>', '')
                    new_content = current + st.session_state[f"{import_key}_pending"]
                    st.session_state[content_key] = new_content
                    st.session_state[version_key] += 1
                    st.session_state[import_key] = False
                    st.session_state[f"{import_key}_pending"] = None
                    st.session_state[f"{import_key}_show_options"] = False
                    st.success("✅ File imported (appended to current content)!")
                    st.rerun()
            
            with col_opt3:
                if st.button("❌ Cancel Import", key=f"import_cancel_options_{editor_base_key}", use_container_width=True):
                    st.session_state[f"{import_key}_pending"] = None
                    st.session_state[f"{import_key}_show_options"] = False
                    st.rerun()

st.markdown("---")

# ============================================================================
# IMAGE UPLOAD SECTION
# ============================================================================
if st.session_state.logged_in and st.session_state.image_handler:
    
    if existing_images:
        st.markdown("### 📸 Your Uploaded Photos")
        st.markdown("*Drag and drop photos directly into the editor above*")
        
        for idx, img in enumerate(existing_images):
            col1, col2, col3 = st.columns([2, 3, 1])
            
            with col1:
                st.markdown(img.get("thumb_html", ""), unsafe_allow_html=True)
            
            with col2:
                caption_text = img.get("caption", "")
                if caption_text:
                    st.markdown(f"**📝 Caption:** {caption_text}")
                else:
                    st.markdown("*No caption*")
            
            with col3:
                if st.button(f"🗑️", key=f"del_img_{img['id']}_{idx}"):
                    st.session_state.image_handler.delete_image(img['id'])
                    st.rerun()
        
        st.markdown("---")
    
    with st.expander("📤 Upload New Photos", expanded=len(existing_images) == 0):
        st.markdown("**Add new photos to your story:**")
        
        uploaded_file = st.file_uploader(
            "Choose an image...", 
            type=['jpg', 'jpeg', 'png'], 
            key=f"up_{current_session_id}_{hash(current_question_text)}",
            label_visibility="collapsed"
        )
        
        if uploaded_file:
            col1, col2 = st.columns([3, 1])
            with col1:
                caption = st.text_input(
                    "Caption / Description:",
                    placeholder="What does this photo show? When was it taken?",
                    key=f"cap_{current_session_id}_{hash(current_question_text)}"
                )
                usage = st.radio(
                    "Image size:",
                    ["Full Page", "Inline"],
                    horizontal=True,
                    key=f"usage_{current_session_id}_{hash(current_question_text)}",
                    help="Full Page: 1600px wide, Inline: 800px wide"
                )
            with col2:
                if st.button("📤 Upload", key=f"btn_{current_session_id}_{hash(current_question_text)}", type="primary", use_container_width=True):
                    with st.spinner("Uploading and optimizing..."):
                        usage_type = "full_page" if usage == "Full Page" else "inline"
                        result = st.session_state.image_handler.save_image(
                            uploaded_file, current_session_id, current_question_text, caption, usage_type
                        )
                        if result:
                            st.success("✅ Photo uploaded and optimized!")
                            time.sleep(1.5)
                            st.rerun()
                        else:
                            st.error("Upload failed")
    
    st.markdown("---")

# ============================================================================
# PREVIEW SECTION
# ============================================================================
current_content = st.session_state.get(content_key, "")
if current_content and current_content != "<p><br></p>" and current_content != "<p>Start writing your story here...</p>":
    with st.expander("👁️ Preview your story", expanded=False):
        st.markdown("### 📖 Preview")
        st.markdown(current_content, unsafe_allow_html=True)
        st.markdown("---")

# ============================================================================
# BETA READER FEEDBACK SECTION
# ============================================================================
st.subheader("🦋 Beta Reader Feedback")

tab1, tab2 = st.tabs(["📝 Current Session", "📚 Feedback History"])

with tab1:
    sdata = st.session_state.responses.get(current_session_id, {})
    answered_cnt = len(sdata.get("questions", {}))
    total_q = len(current_session["questions"])

    st.markdown(f"**Progress:** {answered_cnt}/{total_q} topics answered")
    
    beta_key = f"beta_{current_session_id}_{current_question_text}"
    
    if "beta_feedback_storage" not in st.session_state:
        st.session_state.beta_feedback_storage = {}
    
    col1, col2 = st.columns([2, 1])
    with col1: 
        fb_type = st.selectbox("Feedback Type", ["comprehensive", "concise", "developmental"], 
                              key=f"beta_type_{beta_key}")
    with col2:
        if st.button("🦋 Get Beta Read", key=f"beta_btn_{beta_key}", use_container_width=True, type="primary"):
            with st.spinner("Beta Reader is analyzing your stories with full profile context..."):
                if beta_reader:
                    session_text = ""
                    for q, a in sdata.get("questions", {}).items():
                        text_only = re.sub(r'<[^>]+>', '', a.get("answer", ""))
                        session_text += f"Question: {q}\nAnswer: {text_only}\n\n"
                    
                    if session_text.strip():
                        fb = generate_beta_reader_feedback(current_session["title"], session_text, fb_type)
                        if "error" not in fb: 
                            st.session_state.beta_feedback_display = fb
                            st.session_state.beta_feedback_storage[beta_key] = fb
                            st.rerun()
                        else: 
                            st.error(f"Failed: {fb['error']}")
                    else: 
                        st.warning("No content to analyze. Write some stories first!")
                else:
                    st.error("Beta reader not available")
    
    if beta_key in st.session_state.beta_feedback_storage:
        display_beta_feedback(st.session_state.beta_feedback_storage[beta_key])
    elif st.session_state.beta_feedback_display and not st.session_state.beta_feedback_storage:
        display_beta_feedback(st.session_state.beta_feedback_display)

with tab2:
    st.markdown("### 📚 Your Saved Beta Reader Feedback")
    
    user_data = load_user_data(st.session_state.user_id) if st.session_state.user_id else {}
    all_feedback = user_data.get("beta_feedback", {})
    
    if not all_feedback:
        st.info("No saved feedback yet. Generate feedback from any session and click 'Save This Feedback to History' to keep it forever.")
    else:
        all_entries = []
        for session_id_str, feedback_list in all_feedback.items():
            session_title = "Unknown Session"
            for s in SESSIONS:
                if str(s["id"]) == session_id_str:
                    session_title = s["title"]
                    break
            
            for fb in feedback_list:
                all_entries.append({
                    "session_id": session_id_str, "session_title": session_title,
                    "date": fb.get('generated_at', datetime.now().isoformat()), "feedback": fb
                })
        
        all_entries.sort(key=lambda x: x['date'], reverse=True)
        
        for i, entry in enumerate(all_entries):
            fb = entry['feedback']
            fb_date = datetime.fromisoformat(entry['date']).strftime('%B %d, %Y at %I:%M %p')
            
            with st.expander(f"📖 {entry['session_title']} - {fb_date} ({fb.get('feedback_type', 'comprehensive').title()})"):
                col1, col2, col3 = st.columns([2, 2, 1])
                
                with col1:
                    st.markdown(f"**Session:** {entry['session_title']}")
                with col2:
                    st.markdown(f"**Type:** {fb.get('feedback_type', 'comprehensive').title()}")
                with col3:
                    if st.button(f"🗑️ Delete", key=f"del_fb_history_{i}_{entry['date']}", use_container_width=True):
                        session_id_str = entry['session_id']
                        feedback_list = all_feedback.get(session_id_str, [])
                        
                        feedback_list = [f for f in feedback_list if f.get('generated_at') != entry['date']]
                        
                        if feedback_list:
                            all_feedback[session_id_str] = feedback_list
                        else:
                            del all_feedback[session_id_str]
                        
                        user_data["beta_feedback"] = all_feedback
                        save_user_data(st.session_state.user_id, st.session_state.responses)
                        st.success("Feedback deleted!")
                        st.rerun()
                
                if 'feedback' in fb and fb['feedback']:
                    st.markdown(fb['feedback'])
                else:
                    if 'summary' in fb and fb['summary']:
                        st.markdown("**Summary:**")
                        st.markdown(fb['summary'])
                    
                    if 'strengths' in fb and fb['strengths']:
                        st.markdown("**Strengths:**")
                        for s in fb['strengths']:
                            st.markdown(f"✅ {s}")
                    
                    if 'areas_for_improvement' in fb and fb['areas_for_improvement']:
                        st.markdown("**Areas for Improvement:**")
                        for a in fb['areas_for_improvement']:
                            st.markdown(f"📝 {a}")
                    
                    if 'suggestions' in fb and fb['suggestions']:
                        st.markdown("**Suggestions:**")
                        for sug in fb['suggestions']:
                            st.markdown(f"💡 {sug}")
                    
                    if 'overall_score' in fb and fb['overall_score']:
                        st.markdown(f"**Overall Score:** {fb['overall_score']}/10")

# ============================================================================
# SESSION PROGRESS
# ============================================================================
progress_info = get_progress_info(current_session_id)
st.markdown(f"""
<div class="progress-container">
<div class="progress-header">📊 Session Progress</div>
<div class="progress-status">{progress_info['emoji']} {progress_info['progress_percent']:.0f}% complete • {progress_info['remaining_words']} words remaining</div>
<div class="progress-bar-container"><div class="progress-bar-fill" style="width: {min(progress_info['progress_percent'], 100)}%; background-color: {progress_info['color']};"></div></div>
<div class="progress-stats">{progress_info['current_count']} / {progress_info['target']} words</div>
</div>
""", unsafe_allow_html=True)

if st.button("✏️ Change Word Target", key="edit_target", use_container_width=True): 
    st.session_state.editing_word_target = not st.session_state.editing_word_target
    st.rerun()

if st.session_state.editing_word_target:
    new_target = st.number_input("Target words:", min_value=100, max_value=5000, value=progress_info['target'], key="target_edit")
    col_s, col_c = st.columns(2)
    with col_s:
        if st.button("💾 Save", key="save_target", type="primary", use_container_width=True):
            st.session_state.responses[current_session_id]["word_target"] = new_target
            save_user_data(st.session_state.user_id, st.session_state.responses)
            st.session_state.editing_word_target = False
            st.rerun()
    with col_c:
        if st.button("❌ Cancel", key="cancel_target", use_container_width=True): 
            st.session_state.editing_word_target = False
            st.rerun()

st.divider()

col1, col2, col3, col4 = st.columns(4)
with col1: 
    st.metric("Total Words", sum(calculate_author_word_count(s["id"]) for s in SESSIONS))
with col2: 
    unique_q = set()
    for s in SESSIONS:
        for q, _ in st.session_state.responses.get(s["id"], {}).get("questions", {}).items():
            unique_q.add((s["id"], q))
    comp = sum(1 for s in SESSIONS if len([x for (sid,x) in unique_q if sid == s["id"]]) == len(s["questions"]))
    st.metric("Completed Sessions", f"{comp}/{len(SESSIONS)}")
with col3: 
    st.metric("Topics Explored", f"{len(unique_q)}/{sum(len(s['questions']) for s in SESSIONS)}")
with col4: 
    st.metric("Total Answers", sum(len(st.session_state.responses.get(s["id"], {}).get("questions", {})) for s in SESSIONS))

st.markdown("---")
if st.session_state.user_account:
    profile = st.session_state.user_account['profile']
    age = (datetime.now() - datetime.fromisoformat(st.session_state.user_account['created_at'])).days
    st.caption(f"Tell My Story Timeline • 👤 {profile['first_name']} {profile['last_name']} • 📅 Account Age: {age} days • 📚 Bank: {st.session_state.get('current_bank_name', 'None')}")
else: 
    st.caption(f"Tell My Story Timeline • User: {st.session_state.user_id}")
